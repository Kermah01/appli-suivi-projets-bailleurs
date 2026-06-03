"""
Génération du Cahier des Charges — Application de Suivi des Projets Bailleurs
Ministère de la Planification du Développement — Côte d'Ivoire
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Couleurs ──────────────────────────────────────────────────────────────────
ORANGE        = RGBColor(0xF7, 0x7F, 0x00)   # Orange ministère
DARK_BLUE     = RGBColor(0x1E, 0x3A, 0x5F)   # Bleu foncé
MEDIUM_BLUE   = RGBColor(0x2D, 0x6A, 0xA0)   # Bleu secondaire
LIGHT_GRAY    = RGBColor(0xF2, 0xF4, 0xF7)   # Fond gris clair
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BLACK         = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_TEXT     = RGBColor(0x4B, 0x55, 0x63)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Définit la couleur de fond d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, border_color='BFBFBF'):
    """Ajoute des bordures légères autour d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_page_break(doc):
    doc.add_page_break()


def add_heading(doc, text, level, color=None, space_before=12, space_after=6):
    """Ajoute un titre avec style personnalisé."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    run = h.runs[0] if h.runs else h.add_run(text)
    if color:
        run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        if not color:
            run.font.color.rgb = DARK_BLUE
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        if not color:
            run.font.color.rgb = MEDIUM_BLUE
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        if not color:
            run.font.color.rgb = DARK_BLUE
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=10,
                  color=None, space_after=6, indent=False, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = BLACK
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Ajoute un élément de liste à puces."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = BLACK
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = BLACK
    return p


def add_info_box(doc, text, bg='EBF4FB', border='2D6AA0'):
    """Ajoute un encadré informatif."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, border)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK_BLUE
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_two_col_table(doc, rows_data, header=None, col_widths=None):
    """Table à 2 colonnes avec ligne d'en-tête optionnelle."""
    total = len(rows_data) + (1 if header else 0)
    table = doc.add_table(rows=total, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])

    row_idx = 0
    if header:
        for j, val in enumerate(header):
            cell = table.cell(0, j)
            set_cell_bg(cell, '1E3A5F')
            set_cell_borders(cell, '1E3A5F')
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_idx = 1

    for i, row_data in enumerate(rows_data):
        bg = 'F2F4F7' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = table.cell(row_idx + i, j)
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.clear()
            if isinstance(val, tuple):  # (text, bold)
                run = p.add_run(val[0])
                run.font.bold = val[1]
            else:
                run = p.add_run(str(val))
                run.font.bold = (j == 0)
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_multi_col_table(doc, headers, rows_data, col_widths=None):
    """Table multi-colonnes générique."""
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_bg(cell, '1E3A5F')
        set_cell_borders(cell, '1E3A5F')
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row_data in enumerate(rows_data):
        bg = 'F2F4F7' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = BLACK

    if col_widths:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'F77F00')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── Marges ──
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Police par défaut ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = BLACK

    # ════════════════════════════════════════════════════════════════
    # PAGE DE GARDE
    # ════════════════════════════════════════════════════════════════
    build_cover_page(doc)
    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # SOMMAIRE (manuel)
    # ════════════════════════════════════════════════════════════════
    build_toc(doc)
    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # SECTIONS
    # ════════════════════════════════════════════════════════════════
    build_section_1(doc)
    build_section_2(doc)
    build_section_3(doc)
    build_section_4(doc)
    build_section_5(doc)
    build_section_6(doc)
    build_section_7(doc)
    build_section_8(doc)
    build_section_9(doc)
    build_section_10(doc)

    return doc


# ── Page de garde ─────────────────────────────────────────────────────────────

def build_cover_page(doc):
    # Bandeau haut
    table_header = doc.add_table(rows=1, cols=1)
    table_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table_header.cell(0, 0)
    set_cell_bg(cell, '1E3A5F')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("RÉPUBLIQUE DE CÔTE D'IVOIRE")
    r.font.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = WHITE
    r.font.name = 'Calibri'
    p.add_run('\n')
    r2 = p.add_run("Ministère de la Planification et du Développement")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xBF, 0xD7, 0xEA)
    r2.font.name = 'Calibri'

    for _ in range(4):
        doc.add_paragraph()

    # Titre principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("CAHIER DES CHARGES")
    r_title.font.bold = True
    r_title.font.size = Pt(26)
    r_title.font.color.rgb = DARK_BLUE
    r_title.font.name = 'Calibri'

    # Ligne orange
    table_line = doc.add_table(rows=1, cols=1)
    table_line.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_line = table_line.cell(0, 0)
    set_cell_bg(cell_line, 'F77F00')
    cell_line.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell_line.paragraphs[0].paragraph_format.space_after = Pt(3)
    cell_line.width = Cm(12)

    doc.add_paragraph()

    # Sous-titre
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Application de Suivi des Projets\net des Bailleurs de Fonds")
    r_sub.font.size = Pt(18)
    r_sub.font.bold = True
    r_sub.font.color.rgb = ORANGE
    r_sub.font.name = 'Calibri'

    for _ in range(5):
        doc.add_paragraph()

    # Bloc méta
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta = [
        ("Organisme maître d'ouvrage", "Cabinet du Ministère de la Planification et du Développement — Côte d'Ivoire"),
        ("Type de document",           "Cahier des Charges — Spécifications fonctionnelles & techniques"),
        ("Version",                    "1.0"),
        ("Date de rédaction",          datetime.date.today().strftime("%d %B %Y")),
        ("Statut",                     "Document de référence"),
    ]
    for i, (label, value) in enumerate(meta):
        bg_l = 'EBF4FB' if i % 2 == 0 else 'F2F4F7'
        bg_v = 'FFFFFF' if i % 2 == 0 else 'FAFAFA'
        c0 = table_meta.cell(i, 0)
        c1 = table_meta.cell(i, 1)
        set_cell_bg(c0, bg_l)
        set_cell_bg(c1, bg_v)
        set_cell_borders(c0, 'D1D5DB')
        set_cell_borders(c1, 'D1D5DB')
        p0 = c0.paragraphs[0]
        p0.paragraph_format.left_indent = Cm(0.3)
        p0.paragraph_format.space_before = Pt(5)
        p0.paragraph_format.space_after = Pt(5)
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = DARK_BLUE
        p1 = c1.paragraphs[0]
        p1.paragraph_format.left_indent = Cm(0.3)
        p1.paragraph_format.space_before = Pt(5)
        p1.paragraph_format.space_after = Pt(5)
        r1 = p1.add_run(value)
        r1.font.size = Pt(10)
        r1.font.color.rgb = BLACK

    for _ in range(6):
        doc.add_paragraph()

    # Pied de garde
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("Cabinet du Ministère de la Planification et du Développement  •  Côte d'Ivoire  •  2025")
    r_foot.font.size = Pt(9)
    r_foot.font.color.rgb = GRAY_TEXT
    r_foot.font.italic = True


# ── Sommaire ──────────────────────────────────────────────────────────────────

def build_toc(doc):
    add_heading(doc, "SOMMAIRE", 1, color=DARK_BLUE)
    add_separator(doc)
    doc.add_paragraph()

    toc_entries = [
        ("1.", "Contexte et présentation du projet", 1),
        ("  1.1", "Contexte institutionnel", 2),
        ("  1.2", "Problématique", 2),
        ("  1.3", "Objectifs du projet", 2),
        ("2.", "Périmètre et utilisateurs cibles", 1),
        ("  2.1", "Périmètre applicatif", 2),
        ("  2.2", "Profils utilisateurs", 2),
        ("3.", "Spécifications fonctionnelles détaillées", 1),
        ("  3.1", "Tableau de bord", 2),
        ("  3.2", "Gestion des projets", 2),
        ("  3.3", "Gestion des bailleurs de fonds", 2),
        ("  3.4", "Financements et décaissements", 2),
        ("  3.5", "Plan National de Développement (PND)", 2),
        ("  3.6", "Module d'import de données (Excel)", 2),
        ("  3.7", "Assistant IA (Intelligence Artificielle)", 2),
        ("  3.8", "Gestion des comptes et des accès", 2),
        ("  3.9", "Recherche globale et notifications", 2),
        ("  3.10", "Export des données", 2),
        ("4.", "Architecture technique", 1),
        ("  4.1", "Stack technologique", 2),
        ("  4.2", "Structure applicative (modules Django)", 2),
        ("  4.3", "Modèle de données conceptuel", 2),
        ("  4.4", "Structure des URLs", 2),
        ("5.", "Sécurité et gestion des droits", 1),
        ("  5.1", "Système d'authentification", 2),
        ("  5.2", "Matrice des droits (RBAC)", 2),
        ("  5.3", "Journal d'activité (Audit Log)", 2),
        ("6.", "Interface utilisateur et ergonomie", 1),
        ("7.", "Déploiement et infrastructure", 1),
        ("8.", "Exigences non fonctionnelles", 1),
        ("9.", "Contraintes et dépendances", 1),
        ("10.", "Glossaire", 1),
    ]

    for num, title, lvl in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if lvl == 1:
            p.paragraph_format.left_indent = Cm(0)
        else:
            p.paragraph_format.left_indent = Cm(0.8)
        r_num = p.add_run(f"{num}  ")
        r_num.font.bold = (lvl == 1)
        r_num.font.size = Pt(10 if lvl == 1 else 9.5)
        r_num.font.color.rgb = ORANGE if lvl == 1 else MEDIUM_BLUE
        r_title = p.add_run(title)
        r_title.font.bold = (lvl == 1)
        r_title.font.size = Pt(10 if lvl == 1 else 9.5)
        r_title.font.color.rgb = DARK_BLUE


# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Contexte et présentation
# ════════════════════════════════════════════════════════════════════════════

def build_section_1(doc):
    add_heading(doc, "1.  Contexte et présentation du projet", 1)
    add_separator(doc)

    # 1.1
    add_heading(doc, "1.1  Contexte institutionnel", 2)
    add_paragraph(doc,
        "La Côte d'Ivoire bénéficie d'un soutien financier significatif de la part d'une grande diversité de partenaires "
        "techniques et financiers (PTF) — organismes multilatéraux, bilatéraux, régionaux, ONG internationales et "
        "institutions financières — dans le cadre de la mise en œuvre de son Plan National de Développement (PND). "
        "Ce soutien se traduit par des centaines de projets sectoriels et transversaux couvrant l'ensemble du territoire national.")
    add_paragraph(doc,
        "Le Cabinet du Ministère de la Planification et du Développement est l'institution en charge du pilotage "
        "stratégique, de la coordination, du suivi et de l'évaluation de l'aide publique au développement (APD). "
        "À ce titre, il est le point d'entrée institutionnel de l'ensemble des bailleurs de fonds et assume la "
        "responsabilité de la cohérence entre les interventions des PTF et les priorités nationales de développement.")
    add_info_box(doc,
        "ℹ️  Contexte géographique : L'application est déployée pour les besoins du Ministère de la Planification "
        "et du Développement de la République de Côte d'Ivoire. Les références géographiques (régions, zones) "
        "correspondent au territoire ivoirien. Le fuseau horaire de référence est Africa/Ouagadougou (UTC+0). "
        "La monnaie de référence nationale est le Franc CFA (XOF).")

    # 1.2
    add_heading(doc, "1.2  Problématique", 2)
    add_paragraph(doc,
        "Avant la mise en place de cette solution, le suivi des projets et des bailleurs reposait principalement sur "
        "des outils bureautiques (tableurs Excel, documents Word) dispersés entre différentes directions et points focaux, "
        "rendant la consolidation des données fastidieuse, la traçabilité insuffisante et la production de rapports "
        "chronophage. Les principaux problèmes identifiés étaient :")
    add_bullet(doc, "Absence d'une base de données centralisée et unifiée des projets et des financements.")
    add_bullet(doc, "Difficulté à produire en temps réel des indicateurs clés de performance (KPIs) fiables.")
    add_bullet(doc, "Manque de visibilité sur l'alignement des projets avec les priorités du Plan National de Développement.")
    add_bullet(doc, "Impossibilité de détecter automatiquement les projets en retard ou à faible taux de décaissement.")
    add_bullet(doc, "Absence de mécanisme structuré de gestion des droits d'accès par profil et par bailleur.")

    # 1.3
    add_heading(doc, "1.3  Objectifs du projet", 2)
    add_paragraph(doc,
        "L'application « Suivi des Projets Bailleurs » a été développée pour répondre à l'ensemble de ces enjeux. "
        "Ses objectifs stratégiques sont les suivants :")
    add_bullet(doc, "Centraliser et fiabiliser les données relatives aux projets, bailleurs et financements.",
               bold_prefix="Centralisation — ")
    add_bullet(doc, "Permettre un suivi en temps réel des indicateurs financiers (engagements, décaissements, taux de décaissement) et physiques (avancement).",
               bold_prefix="Suivi en temps réel — ")
    add_bullet(doc, "Mettre en évidence l'alignement des projets sur les piliers et sous-objectifs du Plan National de Développement (PND).",
               bold_prefix="Alignement PND — ")
    add_bullet(doc, "Fournir des outils d'alerte automatique sur les projets en retard et à faible décaissement.",
               bold_prefix="Alerte — ")
    add_bullet(doc, "Offrir une interface d'analyse interactive et des visualisations graphiques avancées pour la prise de décision.",
               bold_prefix="Aide à la décision — ")
    add_bullet(doc, "Garantir la sécurité des données via un système de rôles et de permissions granulaire.",
               bold_prefix="Sécurité — ")
    add_bullet(doc, "Faciliter l'alimentation de la base de données par import en masse depuis des fichiers Excel.",
               bold_prefix="Interopérabilité — ")
    add_bullet(doc, "Intégrer une interface d'interrogation par intelligence artificielle (IA Gemini) pour des analyses conversationnelles.",
               bold_prefix="Intelligence artificielle — ")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Périmètre et utilisateurs
# ════════════════════════════════════════════════════════════════════════════

def build_section_2(doc):
    add_heading(doc, "2.  Périmètre et utilisateurs cibles", 1)
    add_separator(doc)

    # 2.1
    add_heading(doc, "2.1  Périmètre applicatif", 2)
    add_paragraph(doc,
        "L'application couvre l'intégralité du cycle de vie d'un projet de développement financé par un bailleur, "
        "depuis son identification jusqu'à sa clôture. Elle englobe les domaines fonctionnels suivants :")

    modules = [
        ("Tableau de bord analytique", "Vue synthétique des KPIs, graphiques interactifs, alertes en temps réel, carte géographique."),
        ("Gestion des projets", "Référentiel complet des projets : création, consultation, modification, suppression, filtres, export Excel."),
        ("Gestion des bailleurs de fonds", "Répertoire des partenaires financiers avec fiche analytique détaillée par bailleur."),
        ("Financements & décaissements", "Suivi des engagements financiers et des flux de décaissement par projet et par bailleur."),
        ("Plan National de Développement", "Mapping des projets sur les piliers et sous-objectifs du PND actif."),
        ("Import Excel", "Alimentation de la base de données par import de fichiers Excel structurés avec prévisualisation et rapport d'erreurs."),
        ("Assistant IA", "Interface conversationnelle basée sur l'API Gemini de Google pour l'analyse intelligente des données."),
        ("Gestion des comptes", "Authentification, inscription, approbation des comptes, gestion des rôles et journal d'audit."),
    ]
    add_multi_col_table(doc, ["Module", "Description"], modules, col_widths=[5.5, 11])

    # 2.2
    add_heading(doc, "2.2  Profils utilisateurs", 2)
    add_paragraph(doc,
        "L'application définit quatre profils distincts avec des niveaux d'accès progressifs. "
        "Chaque compte doit être approuvé par un administrateur avant d'être activé.")

    profils = [
        ("Super Administrateur", "superadmin",
         "Accès total à toutes les données et fonctionnalités. Gère les comptes, les rôles et les données de toutes structures."),
        ("Directeur / Haute Fonction", "directeur",
         "Accès complet en lecture et écriture sur toutes les données. Destiné au Directeur de Cabinet, Chef de Cabinet, Conseillers et Directeurs Généraux."),
        ("Point Focal", "point_focal",
         "Accès restreint aux bailleurs qui lui sont explicitement assignés. Peut créer et modifier des projets et financements liés à ses bailleurs uniquement."),
        ("Lecteur", "lecteur",
         "Accès en lecture seule sur les données visibles selon ses bailleurs assignés. Aucune modification possible."),
    ]
    headers = ["Profil", "Code interne", "Description et droits"]
    add_multi_col_table(doc, headers,
        [(p[0], p[1], p[2]) for p in profils],
        col_widths=[4, 3, 10])

    add_paragraph(doc,
        "Les fonctions institutionnelles couvertes incluent : Ministre, Directeur de Cabinet, Directeur de Cabinet Adjoint, "
        "Chef de Cabinet, Conseiller Technique, Directeur Général, Chargé d'Études et Point Focal. "
        "Les fonctions à responsabilité unique (Ministre, Directeur de Cabinet, etc.) ne peuvent être attribuées "
        "qu'à un seul compte approuvé à la fois.")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Spécifications fonctionnelles
# ════════════════════════════════════════════════════════════════════════════

def build_section_3(doc):
    add_heading(doc, "3.  Spécifications fonctionnelles détaillées", 1)
    add_separator(doc)

    # 3.1 — Tableau de bord
    add_heading(doc, "3.1  Tableau de bord analytique", 2)
    add_paragraph(doc,
        "Le tableau de bord est la page d'accueil principale de l'application. Il offre une vision consolidée "
        "et en temps réel de l'ensemble du portefeuille de projets et des financements. Son contenu est "
        "automatiquement filtré selon le profil de l'utilisateur connecté : un Point Focal ne voit que les "
        "données relatives à ses bailleurs assignés.")

    add_heading(doc, "3.1.1  Indicateurs clés de performance (KPIs)", 3)
    add_paragraph(doc,
        "Quatre cartes KPI sont affichées en haut du tableau de bord, avec une variation calculée "
        "sur les 30 derniers jours (flèche haut/bas colorée) :")
    kpis = [
        ("Nombre total de projets", "Tous statuts confondus dans le périmètre de l'utilisateur."),
        ("Projets en cours d'exécution", "Projets avec le statut « En cours d'exécution »."),
        ("Projets en retard", "Projets en cours dont la date de fin prévue est dépassée."),
        ("Nombre de bailleurs actifs", "Bailleurs ayant au moins un financement dans le périmètre."),
    ]
    add_two_col_table(doc, kpis, header=["KPI", "Définition"])

    add_paragraph(doc,
        "Deux indicateurs financiers complémentaires sont affichés sous forme de jauges : "
        "le montant total engagé (cumul de tous les financements) et le montant total décaissé "
        "(cumul de tous les décaissements), ainsi que le taux de décaissement global (%) représenté "
        "par un graphique en arcs radialBar.")

    add_heading(doc, "3.1.2  Graphiques de synthèse", 3)
    charts = [
        ("Répartition par statut", "Graphique en barres horizontales — nombre de projets par statut (Identification, Préparation, Négociation, En cours, Suspendu, Clôturé, Annulé)."),
        ("Répartition par secteur", "Graphique en donut — nombre de projets par secteur d'activité."),
        ("Carte des régions", "Carte interactive de la Côte d'Ivoire (Leaflet.js + GeoJSON) avec localisation des projets par zone géographique."),
        ("Couverture du PND", "Tableau de bord de la couverture du Plan National de Développement actif : nombre de projets et montant engagé par pilier."),
    ]
    add_two_col_table(doc, charts, header=["Graphique", "Description"])

    add_heading(doc, "3.1.3  Moteur d'analyse interactive", 3)
    add_paragraph(doc,
        "Un moteur d'analyse en JavaScript (Alpine.js + ApexCharts) permet à l'utilisateur de générer "
        "des visualisations personnalisées sans rechargement de page. Il dispose de :")
    add_bullet(doc, "Filtres dynamiques : statut, secteur, bailleur, catégorie institutionnelle, zone géographique, co-financement.")
    add_bullet(doc, "Choix du type de graphique : barres, barres horizontales, courbes, aires, donut, camembert, radar, treemap, carte chaleur.")
    add_bullet(doc, "Choix des axes X et Y parmi une liste d'indicateurs : secteur, bailleur, statut, zone, taux d'avancement, taux de décaissement, montant.")
    add_bullet(doc, "Résultats affichés instantanément sur la base des données chargées côté client (JSON).")

    add_heading(doc, "3.1.4  Tableaux de suivi rapide", 3)
    add_bullet(doc, "Tableau des 5 derniers projets créés avec lien vers leur fiche détaillée.")
    add_bullet(doc, "Tableau des projets en retard (5 premiers par ancienneté de retard) avec bailleur et date d'échéance dépassée.")

    # 3.2 — Projets
    add_heading(doc, "3.2  Gestion des projets", 2)
    add_paragraph(doc,
        "Le module Projets est le cœur de l'application. Il gère le référentiel complet de tous les projets "
        "de développement financés par les bailleurs.")

    add_heading(doc, "3.2.1  Données d'un projet", 3)
    fields = [
        ("Code du projet", "Identifiant unique alphanumérique (obligatoire)."),
        ("Titre", "Intitulé complet du projet (obligatoire)."),
        ("Description", "Texte libre de description."),
        ("Secteur", "Lien vers le référentiel des secteurs (Agriculture, Éducation, Santé, etc.)."),
        ("Bailleur principal", "Référence au bailleur principal du projet."),
        ("Statut", "Identification / Préparation / Négociation / En cours / Suspendu / Clôturé / Annulé."),
        ("Montant total", "Montant global du projet dans la devise renseignée."),
        ("Devise", "USD, EUR, XOF, GBP, JPY, CHF, CNY."),
        ("Taux d'avancement physique", "Pourcentage d'exécution physique (0–100 %)."),
        ("Date de signature", "Date de signature de la convention de financement."),
        ("Date de début", "Date de démarrage effectif du projet."),
        ("Date de fin prévue", "Date de clôture contractuelle prévisionnelle."),
        ("Date de fin effective", "Date de clôture réelle (renseignée après clôture)."),
        ("Zone géographique", "Région(s) d'intervention sur le territoire ivoirien."),
        ("Responsable / Chef de projet", "Nom du responsable côté gouvernement."),
        ("Email et téléphone du responsable", "Coordonnées du responsable."),
        ("Structure responsable", "Ministère ou direction en charge de l'exécution."),
        ("Motif du retard", "Explication textuelle des causes de retard éventuel."),
        ("Objectifs PND", "Liaison Many-to-Many avec les sous-objectifs du Plan National de Développement."),
    ]
    add_multi_col_table(doc, ["Champ", "Description"], fields, col_widths=[5, 11.5])

    add_heading(doc, "3.2.2  Fonctionnalités CRUD", 3)
    add_bullet(doc, "Création d'un projet avec gestion des financements associés directement dans le formulaire (Alpine.js).")
    add_bullet(doc, "Consultation de la fiche détaillée : informations générales, financements, décaissements, alignement PND.")
    add_bullet(doc, "Modification complète des données du projet.")
    add_bullet(doc, "Suppression avec confirmation (réservée aux utilisateurs autorisés).")

    add_heading(doc, "3.2.3  Liste et filtres", 3)
    add_bullet(doc, "Recherche textuelle par code ou titre.")
    add_bullet(doc, "Filtres par statut, secteur, bailleur.")
    add_bullet(doc, "Filtre « projets en retard » (case à cocher).")
    add_bullet(doc, "Indicateurs calculés à la volée : taux de décaissement, statut d'avancement, co-financement.")

    add_heading(doc, "3.2.4  Indicateurs calculés automatiquement", 3)
    indicators = [
        ("Taux de décaissement", "Total décaissé / Montant total × 100 — calculé depuis les décaissements réels."),
        ("Total engagé", "Somme des montants engagés dans tous les financements du projet."),
        ("Total décaissé", "Somme de tous les montants de décaissements liés au projet."),
        ("Nombre de bailleurs", "Nombre de bailleurs distincts finançant le projet (co-financement si > 1)."),
        ("Montant en FCFA", "Conversion automatique dans la devise nationale selon taux indicatifs."),
        ("Statut de retard", "Calculé dynamiquement : retard si statut='en_cours' et date_fin_prevue < aujourd'hui."),
    ]
    add_two_col_table(doc, indicators, header=["Indicateur", "Méthode de calcul"])

    add_heading(doc, "3.2.5  Export Excel", 3)
    add_paragraph(doc,
        "Un bouton d'export permet de télécharger la liste des projets (avec les filtres actifs) au format Excel (.xlsx). "
        "Le fichier exporté comprend : Code, Titre, Bailleur principal, Secteur, Statut, Montant total, Devise, "
        "Taux d'avancement, Zone géographique, Responsable, Structure responsable, Dates, Montant décaissé, "
        "Taux de décaissement et Motif de retard. L'en-tête est formaté avec une couleur institutionnelle "
        "et les lignes sont alternées pour la lisibilité.")

    # 3.3 — Bailleurs
    add_heading(doc, "3.3  Gestion des bailleurs de fonds", 2)
    add_paragraph(doc,
        "Le module Bailleurs constitue le répertoire complet de tous les partenaires financiers intervenant "
        "dans le portefeuille de projets ivoirien.")

    add_heading(doc, "3.3.1  Données d'un bailleur", 3)
    bailleur_fields = [
        ("Nom complet", "Dénomination officielle complète de l'institution."),
        ("Sigle / Acronyme", "Identifiant court (ex. BM, AFD, BAfD, PNUD)."),
        ("Type de bailleur", "Multilatéral / Bilatéral / Régional / Privé / ONG Internationale / Autre."),
        ("Catégorie institutionnelle", "Classification fine : Bretton Woods, Système ONU, Banques multilatérales, Coopération bilatérale, Institutions régionales africaines, Fonds verticaux/thématiques, Secteur privé/Fondations, ONG internationales."),
        ("Pays du siège", "Localisation du siège social."),
        ("Description", "Texte de présentation de l'institution."),
        ("Site web", "URL officiel."),
        ("Email de contact", "Adresse email de contact institutionnel."),
    ]
    add_multi_col_table(doc, ["Champ", "Description"], bailleur_fields, col_widths=[5, 11.5])

    add_heading(doc, "3.3.2  Fiche analytique bailleur", 3)
    add_paragraph(doc,
        "La page de détail d'un bailleur affiche un tableau de bord analytique complet comprenant :")
    add_bullet(doc, "KPIs : nombre de projets financés, montant total engagé, montant total décaissé, taux de décaissement.")
    add_bullet(doc, "Graphique de répartition des projets par secteur (donut ou barres).")
    add_bullet(doc, "Graphique de répartition par statut des projets.")
    add_bullet(doc, "Graphique de répartition par zone géographique.")
    add_bullet(doc, "Graphique de répartition par type de financement (don, prêt, assistance technique, etc.).")
    add_bullet(doc, "Carte Leaflet des projets localisés dans les régions de Côte d'Ivoire.")
    add_bullet(doc, "Tableau détaillé des projets financés avec indicateurs d'avancement et de décaissement.")
    add_bullet(doc, "Tableau des financements avec montants engagés, décaissés, taux et reste à décaisser.")

    # 3.4 — Financements
    add_heading(doc, "3.4  Financements et décaissements", 2)
    add_paragraph(doc,
        "Le module Financements assure le suivi détaillé des flux financiers entre les bailleurs et les projets.")

    add_heading(doc, "3.4.1  Financement", 3)
    fin_fields = [
        ("Projet", "Référence au projet financé."),
        ("Bailleur", "Référence au bailleur pourvoyeur de fonds."),
        ("Type de financement", "Don / Prêt concessionnel / Prêt non concessionnel / Assistance technique / Cofinancement / Contrepartie nationale / Autre."),
        ("Montant engagé", "Montant du financement accordé dans la devise de l'accord."),
        ("Devise", "USD, EUR, XOF, GBP, JPY, CHF."),
        ("Date d'accord", "Date de signature ou d'entrée en vigueur de l'accord."),
        ("Référence accord", "Numéro ou code de référence de l'accord de financement."),
        ("Observations", "Notes complémentaires."),
    ]
    add_multi_col_table(doc, ["Champ", "Description"], fin_fields, col_widths=[5, 11.5])

    add_heading(doc, "3.4.2  Décaissement", 3)
    dec_fields = [
        ("Financement", "Référence au financement parent (lien bailleur-projet)."),
        ("Montant décaissé", "Montant de ce versement spécifique."),
        ("Date de décaissement", "Date du versement effectif (obligatoire)."),
        ("Référence", "Code de traçabilité du virement."),
        ("Description", "Libellé ou commentaire du décaissement."),
    ]
    add_multi_col_table(doc, ["Champ", "Description"], dec_fields, col_widths=[5, 11.5])

    add_heading(doc, "3.4.3  Indicateurs calculés", 3)
    calc_fields = [
        ("Total décaissé (par financement)", "Somme de tous les décaissements liés à un financement."),
        ("Taux de décaissement (par financement)", "Total décaissé / Montant engagé × 100."),
        ("Reste à décaisser", "Montant engagé – Total décaissé."),
        ("Taux de décaissement global", "Décaissements totaux / Engagements totaux × 100, affiché sur le tableau de bord."),
    ]
    add_two_col_table(doc, calc_fields, header=["Indicateur", "Formule"])

    # 3.5 — PND
    add_heading(doc, "3.5  Plan National de Développement (PND)", 2)
    add_paragraph(doc,
        "Le module PND permet d'aligner les projets sur la structure stratégique nationale et d'identifier "
        "les gaps de financement par axe prioritaire.")

    add_heading(doc, "3.5.1  Structure du PND", 3)
    pnd_struct = [
        ("Plan National de Développement", "Entité racine (ex. PND 2021-2025) avec période de validité et statut actif."),
        ("Pilier", "Axe stratégique du plan (numéroté). Plusieurs piliers par plan."),
        ("Sous-objectif", "Objectif spécifique rattaché à un pilier (numéroté, intitulé)."),
    ]
    add_two_col_table(doc, pnd_struct, header=["Niveau", "Description"])

    add_heading(doc, "3.5.2  Alignement projets-PND", 3)
    add_paragraph(doc,
        "Chaque projet peut être lié à plusieurs sous-objectifs du PND (relation Many-to-Many). "
        "Le module PND affiche pour chaque pilier du plan actif :")
    add_bullet(doc, "Le nombre de projets couvrant ce pilier (directement ou via ses sous-objectifs).")
    add_bullet(doc, "Le montant total des financements engagés sur ce pilier.")
    add_bullet(doc, "La liste des sous-objectifs avec leur nombre de projets associés.")
    add_bullet(doc, "La fiche détaillée d'un pilier liste tous les projets alignés avec leurs indicateurs financiers et physiques.")

    # 3.6 — Import Excel
    add_heading(doc, "3.6  Module d'import de données (Excel)", 2)
    add_paragraph(doc,
        "Ce module permet d'alimenter la base de données en masse depuis un fichier Excel structuré, "
        "sans compétence technique requise. Il est conçu pour les phases d'initialisation ou de mise à jour périodique.")

    add_heading(doc, "3.6.1  Structure du fichier d'import", 3)
    add_paragraph(doc,
        "Le fichier Excel doit contenir jusqu'à quatre feuilles dans l'ordre suivant (chaque feuille est optionnelle) :")
    sheets = [
        ("Bailleurs", "Sigle*, Nom complet*, Type de bailleur, Catégorie institutionnelle, Pays du siège, Description, Site web, Email de contact."),
        ("Projets", "Code projet*, Titre*, Bailleur principal (sigle), Secteur (code), Statut, Montant total, Devise, Taux d'avancement, Date de signature, Date de début, Date de fin prévue, Zone géographique, Responsable."),
        ("Financements", "Code projet*, Sigle bailleur*, Type de financement, Montant engagé, Devise, Date d'accord, Référence accord, Observations."),
        ("Décaissements", "Code projet*, Sigle bailleur*, Montant décaissé*, Date de décaissement*, Référence, Description."),
    ]
    add_two_col_table(doc, sheets, header=["Feuille", "Colonnes (* = obligatoire)"])

    add_heading(doc, "3.6.2  Processus d'import", 3)
    add_paragraph(doc, "Le processus se déroule en deux étapes :")
    add_bullet(doc, "Étape 1 — Prévisualisation (Analyse) : L'utilisateur téléverse le fichier. Le moteur analyse chaque ligne et génère un rapport de prévisualisation indiquant les créations prévues, les mises à jour, et les erreurs détectées (champs manquants, références introuvables, valeurs invalides).",
               bold_prefix="Prévisualisation — ")
    add_bullet(doc, "Étape 2 — Confirmation (Import) : Après validation de la prévisualisation, l'utilisateur confirme l'import. Le moteur exécute les opérations en base dans l'ordre : Bailleurs → Projets → Financements → Décaissements. Un rapport final résume le nombre d'enregistrements créés, mis à jour et les erreurs.",
               bold_prefix="Exécution — ")
    add_paragraph(doc,
        "La logique d'import est idempotente : une ligne déjà existante est mise à jour (update_or_create) "
        "plutôt que dupliquée, sur la base d'identifiants naturels (sigle pour les bailleurs, code pour les projets, "
        "combinaison projet+bailleur+type pour les financements, référence+date pour les décaissements).")

    # 3.7 — Assistant IA
    add_heading(doc, "3.7  Assistant IA (Intelligence Artificielle)", 2)
    add_paragraph(doc,
        "L'application intègre un assistant conversationnel basé sur l'API Google Gemini, accessible depuis "
        "le menu principal. Il permet aux utilisateurs d'interroger la base de données en langage naturel "
        "et d'obtenir des réponses structurées, enrichies de graphiques et de tableaux.")

    add_heading(doc, "3.7.1  Fonctionnement", 3)
    add_paragraph(doc,
        "À chaque requête, l'assistant reçoit un contexte JSON complet de la base de données (statistiques globales, "
        "détail des projets, bailleurs, financements, décaissements) filtré selon le profil de l'utilisateur. "
        "Ce contexte est enrichi d'une instruction système définissant le rôle et le format de réponse attendu.")
    add_paragraph(doc,
        "L'assistant est capable de produire :")
    add_bullet(doc, "Des synthèses textuelles en Markdown (gras, tableaux, listes).")
    add_bullet(doc, "Des graphiques dynamiques (ApexCharts) : barres, courbes, donut, camembert, aires, radialBar, treemap.")
    add_bullet(doc, "Des tableaux de données structurés.")

    add_heading(doc, "3.7.2  Modèles IA utilisés", 3)
    add_paragraph(doc,
        "L'application tente les modèles suivants dans l'ordre, en cas de quota dépassé : "
        "gemini-2.5-flash → gemini-2.0-flash → gemini-2.5-flash-lite → gemini-2.0-flash-lite. "
        "L'historique de conversation est transmis à chaque échange pour maintenir la cohérence du dialogue. "
        "La clé API est configurée via la variable d'environnement GEMINI_API_KEY.")

    add_heading(doc, "3.7.3  Filtrage des données par profil", 3)
    add_paragraph(doc,
        "Le contexte transmis à l'IA est strictement filtré selon les mêmes règles de droits que l'interface : "
        "un Point Focal ne verra dans les réponses de l'IA que les données de ses bailleurs assignés. "
        "Cela garantit la confidentialité et la pertinence des analyses.")

    # 3.8 — Comptes
    add_heading(doc, "3.8  Gestion des comptes et des accès", 2)
    add_paragraph(doc, "Ce module gère l'ensemble du cycle de vie des utilisateurs de la plateforme.")
    add_bullet(doc, "Page d'inscription publique avec saisie des informations personnelles et professionnelles (nom, prénom, email, fonction, téléphone). Les comptes créés sont désactivés par défaut jusqu'à approbation.")
    add_bullet(doc, "Page de connexion avec redirection automatique après authentification réussie.")
    add_bullet(doc, "Panneau d'administration des comptes (super admins uniquement) : liste filtrée des profils, approbation/rejet/suspension, changement de rôle, édition complète, assignation des bailleurs.")
    add_bullet(doc, "Journal d'audit consultable (200 dernières entrées) : création, modification, suppression, approbation, connexion.")
    add_paragraph(doc,
        "Lors de l'inscription, les fonctions à responsabilité unique (Ministre, Directeur de Cabinet, etc.) "
        "sont automatiquement grisées si elles sont déjà occupées par un compte approuvé. "
        "Un message d'attente est présenté après l'inscription, avec possibilité pour l'administrateur "
        "d'approuver ou rejeter le compte depuis l'interface de gestion.")

    # 3.9 — Recherche et notifications
    add_heading(doc, "3.9  Recherche globale et notifications", 2)
    add_paragraph(doc,
        "Une barre de recherche globale accessible depuis le menu principal permet de rechercher "
        "simultanément des projets (par code ou titre), des bailleurs (par nom ou sigle) et "
        "des financements (par code projet, sigle bailleur ou référence). Les résultats sont affichés "
        "en temps réel via une API JSON sans rechargement de page (minimum 2 caractères saisis).")
    add_paragraph(doc,
        "Un centre de notifications accessible via une icône cloche dans la barre de navigation "
        "présente les alertes actives :")
    add_bullet(doc, "Projets en retard (avec nombre de jours de dépassement).")
    add_bullet(doc, "Projets en cours à faible taux de décaissement (< 20 %).")
    add_bullet(doc, "Activités récentes (créations/modifications des 7 derniers jours).")
    add_paragraph(doc,
        "Le contenu de ces notifications est filtré par profil utilisateur. Le badge du compteur "
        "est mis à jour automatiquement à chaque ouverture du panneau.")

    # 3.10 — Export
    add_heading(doc, "3.10  Export des données", 2)
    add_paragraph(doc,
        "L'application offre la possibilité d'exporter les données pour un travail hors connexion :")
    add_bullet(doc, "Export Excel des projets (avec les filtres actifs) : fichier .xlsx mis en forme, avec en-têtes colorées, alternance de lignes et largeurs de colonnes optimisées.")
    add_bullet(doc, "Le nom du fichier exporté inclut la date du jour (ex. projets_20250511.xlsx).")
    add_bullet(doc, "Capacité à étendre l'export aux bailleurs et aux financements dans des versions futures.")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Architecture technique
# ════════════════════════════════════════════════════════════════════════════

def build_section_4(doc):
    add_heading(doc, "4.  Architecture technique", 1)
    add_separator(doc)

    # 4.1
    add_heading(doc, "4.1  Stack technologique", 2)
    stack = [
        ("Langage principal", "Python 3.11"),
        ("Framework web", "Django 5.2 (LTS)"),
        ("Base de données (dev)", "SQLite (fichier local db.sqlite3)"),
        ("Base de données (prod)", "PostgreSQL (via dj-database-url)"),
        ("Serveur WSGI", "Gunicorn 21.2+"),
        ("Gestion des fichiers statiques", "WhiteNoise 6.6+ (compression + cache)"),
        ("Framework CSS", "TailwindCSS (CDN)"),
        ("Composants réactifs", "Alpine.js (CDN)"),
        ("Graphiques interactifs", "ApexCharts v3.44.0 (CDN)"),
        ("Cartographie", "Leaflet.js (CDN) + GeoJSON Côte d'Ivoire"),
        ("Icônes", "Material Symbols (Google Fonts CDN)"),
        ("Manipulation Excel", "openpyxl 3.1+"),
        ("Connecteur PostgreSQL", "psycopg2-binary 2.9+"),
        ("Intelligence artificielle", "Google Generative AI (Gemini) — google-generativeai 0.3+"),
        ("Plateforme de déploiement", "Render.com (cloud PaaS)"),
        ("Langue de l'interface", "Français (fr-fr), fuseau Africa/Ouagadougou"),
    ]
    add_two_col_table(doc, stack, header=["Composant", "Technologie / Version"])

    # 4.2
    add_heading(doc, "4.2  Structure applicative (modules Django)", 2)
    apps = [
        ("config", "Configuration globale du projet Django (settings.py, urls.py, wsgi.py, asgi.py)."),
        ("dashboard", "Vue principale, calcul des KPIs, APIs JSON (recherche, notifications, GeoJSON), moteur de graphiques."),
        ("bailleurs", "Modèle Bailleur, vues CRUD, fiche analytique bailleur."),
        ("projets", "Modèles Projet et Secteur, vues CRUD, export Excel, taux calculés."),
        ("financements", "Modèles Financement et Decaissement, vues CRUD, indicateurs de décaissement."),
        ("pnd", "Modèles PlanNational, Pilier, SousObjectif, vues couverture PND."),
        ("accounts", "Modèles UserProfile et ActivityLog, vues authentification, gestion des comptes, décorateurs de permission."),
        ("imports", "Moteur d'import Excel (analyse + exécution), vues upload/prévisualisation/confirmation."),
        ("assistant", "Service Gemini (construction du contexte DB + appel API), vues interface de chat."),
    ]
    add_multi_col_table(doc, ["Module", "Rôle"], apps, col_widths=[4, 12.5])

    # 4.3
    add_heading(doc, "4.3  Modèle de données conceptuel", 2)
    add_paragraph(doc,
        "Le schéma de données s'articule autour de cinq entités principales, reliées par des clés étrangères "
        "et une relation Many-to-Many :")

    entities = [
        ("Bailleur", "nom, sigle, type_bailleur, categorie_institutionnelle, pays_siege, description, site_web, contact_email"),
        ("Secteur", "nom, code, description, couleur (hex)"),
        ("Projet", "code (unique), titre, description, secteur (FK→Secteur), bailleur_principal (FK→Bailleur), statut, montant_total, devise, taux_avancement, dates (signature/début/fin_prévue/fin_effective), zone_geographique, responsable, motif_retard, objectifs_pnd (M2M→SousObjectif)"),
        ("Financement", "projet (FK→Projet), bailleur (FK→Bailleur), type_financement, montant_engage, devise, date_accord, reference, observations"),
        ("Decaissement", "financement (FK→Financement), montant, date_decaissement, reference, description"),
        ("PlanNational", "nom, sigle, annee_debut, annee_fin, actif"),
        ("Pilier", "plan (FK→PlanNational), numero, nom, description"),
        ("SousObjectif", "pilier (FK→Pilier), numero, nom, description — relié aux Projets via M2M"),
        ("UserProfile", "user (1:1→User), role, fonction, titre_poste, bailleurs (M2M→Bailleur), is_approved"),
        ("ActivityLog", "user (FK→User), action, model_name, object_repr, timestamp"),
    ]
    add_multi_col_table(doc, ["Entité", "Principaux champs"], entities, col_widths=[4, 12.5])

    add_heading(doc, "4.3.1  Relations clés", 3)
    add_bullet(doc, "Un Projet a un bailleur_principal et peut avoir plusieurs financements issus de bailleurs différents (co-financement).")
    add_bullet(doc, "Un Financement lie un Bailleur à un Projet et peut avoir plusieurs Décaissements.")
    add_bullet(doc, "Un Projet est lié à zéro ou plusieurs SousObjectifs du PND (relation M2M).")
    add_bullet(doc, "Un UserProfile peut être assigné à plusieurs Bailleurs (pour les Points Focaux).")
    add_bullet(doc, "Les taux de décaissement sont tous calculés dynamiquement (propriétés Python), non stockés en base.")

    add_heading(doc, "4.3.2  Devises et conversions", 3)
    add_paragraph(doc,
        "L'application gère 7 devises : USD, EUR, XOF, GBP, JPY, CHF, CNY. "
        "Des taux de conversion indicatifs vers le Franc CFA (XOF) sont intégrés en dur dans le code "
        "(fichier projets/models.py) pour permettre des comparaisons homogènes dans les analyses. "
        "Ces taux sont à mettre à jour périodiquement par un administrateur technique.")

    # 4.4
    add_heading(doc, "4.4  Structure des URLs", 2)
    urls = [
        ("/", "Tableau de bord principal"),
        ("/projets/", "Liste des projets"),
        ("/projets/<id>/", "Fiche détaillée d'un projet"),
        ("/projets/nouveau/", "Formulaire de création d'un projet"),
        ("/projets/<id>/modifier/", "Formulaire de modification"),
        ("/projets/<id>/supprimer/", "Suppression d'un projet"),
        ("/projets/exporter/", "Export Excel des projets"),
        ("/bailleurs/", "Liste des bailleurs"),
        ("/bailleurs/<id>/", "Fiche détaillée d'un bailleur"),
        ("/bailleurs/nouveau/", "Création d'un bailleur"),
        ("/financements/", "Liste des financements"),
        ("/financements/<id>/", "Détail d'un financement"),
        ("/pnd/", "Vue couverture PND"),
        ("/pnd/piliers/<id>/", "Détail d'un pilier PND"),
        ("/import/", "Interface d'import Excel"),
        ("/assistant/", "Interface assistant IA"),
        ("/comptes/connexion/", "Page de connexion"),
        ("/comptes/inscription/", "Page d'inscription"),
        ("/comptes/utilisateurs/", "Gestion des comptes (admin)"),
        ("/comptes/audit/", "Journal d'audit"),
        ("/api/search/", "API JSON — Recherche globale"),
        ("/api/notifications/", "API JSON — Notifications"),
        ("/api/regions-geojson/", "API JSON — GeoJSON régions CI"),
        ("/admin/", "Interface d'administration Django"),
    ]
    add_multi_col_table(doc, ["URL", "Description"], urls, col_widths=[6.5, 10])


# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Sécurité et droits
# ════════════════════════════════════════════════════════════════════════════

def build_section_5(doc):
    add_heading(doc, "5.  Sécurité et gestion des droits", 1)
    add_separator(doc)

    # 5.1
    add_heading(doc, "5.1  Système d'authentification", 2)
    add_paragraph(doc,
        "L'application utilise le système d'authentification natif de Django, complété par un système "
        "d'approbation obligatoire des comptes :")
    add_bullet(doc, "Toutes les pages sont protégées par un décorateur personnalisé login_required_custom qui redirige vers la page de connexion si l'utilisateur n'est pas authentifié.")
    add_bullet(doc, "Un mécanisme supplémentaire vérifie que le compte est approuvé (is_approved=True). Un compte authentifié mais non approuvé est redirigé vers une page d'attente.")
    add_bullet(doc, "Les comptes superuser Django (créés via manage.py createsuperuser) sont exonérés de l'obligation d'approbation.")
    add_bullet(doc, "Les mots de passe sont soumis aux validateurs Django standard : similarité avec l'identifiant, longueur minimale, liste des mots de passe courants, interdiction des mots de passe purement numériques.")
    add_bullet(doc, "La protection CSRF est active sur toutes les requêtes POST.")

    # 5.2
    add_heading(doc, "5.2  Matrice des droits (RBAC)", 2)
    add_paragraph(doc,
        "Le système de contrôle d'accès basé sur les rôles (Role-Based Access Control) s'applique "
        "de manière transversale à toutes les vues et opérations :")

    matrix = [
        ("Consulter le tableau de bord", "✓ (périmètre global)", "✓ (périmètre global)", "✓ (périmètre restreint)", "✓ (périmètre restreint)"),
        ("Voir la liste des projets", "✓ Tous", "✓ Tous", "✓ Ses bailleurs uniquement", "✓ Ses bailleurs uniquement"),
        ("Créer / Modifier un projet", "✓", "✓", "✓ (ses bailleurs uniquement)", "✗"),
        ("Supprimer un projet", "✓", "✓", "✓ (ses bailleurs uniquement)", "✗"),
        ("Créer / Modifier un bailleur", "✓", "✓", "✓ (ses bailleurs uniquement)", "✗"),
        ("Importer des données Excel", "✓", "✓", "✗", "✗"),
        ("Accéder à l'assistant IA", "✓", "✓", "✓", "✓"),
        ("Gérer les comptes utilisateurs", "✓", "✗", "✗", "✗"),
        ("Consulter le journal d'audit", "✓", "✗", "✗", "✗"),
        ("Approuver / rejeter des comptes", "✓", "✗", "✗", "✗"),
    ]
    headers_m = ["Action", "Super Admin", "Directeur", "Point Focal", "Lecteur"]
    add_multi_col_table(doc, headers_m, matrix, col_widths=[5.5, 2.5, 2.5, 2.8, 2.5])

    add_paragraph(doc,
        "Le décorateur edit_permission_required applique un contrôle supplémentaire au niveau de l'objet : "
        "pour les Points Focaux, la modification d'un projet ou d'un financement n'est autorisée que si "
        "le bailleur concerné fait partie de leurs bailleurs assignés. "
        "Un message d'erreur explicite est affiché en cas de tentative non autorisée.")

    # 5.3
    add_heading(doc, "5.3  Journal d'activité (Audit Log)", 2)
    add_paragraph(doc,
        "Toutes les opérations sensibles sont enregistrées automatiquement dans la table ActivityLog :")
    add_bullet(doc, "Actions tracées : création, modification, suppression, approbation de compte, connexion.")
    add_bullet(doc, "Chaque entrée consigne : l'utilisateur, le type d'action, l'objet concerné (modèle + représentation + identifiant), les détails et l'horodatage.")
    add_bullet(doc, "Le journal est consultable par les Super Administrateurs depuis l'interface /comptes/audit/.")
    add_bullet(doc, "Les 3 actions les plus récentes du journal apparaissent également dans le centre de notifications.")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Interface utilisateur
# ════════════════════════════════════════════════════════════════════════════

def build_section_6(doc):
    add_heading(doc, "6.  Interface utilisateur et ergonomie", 1)
    add_separator(doc)

    add_paragraph(doc,
        "L'interface est conçue pour être utilisée depuis un navigateur web moderne (Chrome, Firefox, Edge) "
        "sur un écran d'ordinateur. Elle adopte une approche Single Page Application (SPA) légère via Alpine.js "
        "pour les composants interactifs, sans framework JavaScript lourd.")

    add_heading(doc, "6.1  Charte graphique", 2)
    add_bullet(doc, "Couleur principale : Bleu foncé institutionnel (#1E3A5F) pour les en-têtes et éléments de navigation.")
    add_bullet(doc, "Couleur d'accent : Orange (#F77F00) pour les boutons d'action, badges et éléments de mise en valeur.")
    add_bullet(doc, "Typographie : Calibri / Inter / Système — lisible sur tous les écrans.")
    add_bullet(doc, "Icônes : Material Symbols (Google) — cohérence visuelle sur l'ensemble de l'application.")
    add_bullet(doc, "Design : Interface épurée avec fond clair, cartes avec ombres légères, tableaux alternés.")

    add_heading(doc, "6.2  Navigation principale", 2)
    add_paragraph(doc,
        "Une barre latérale fixe (sidebar) affiche les liens de navigation vers tous les modules, "
        "avec badge de notification actif. Une barre supérieure (topbar) affiche le profil de l'utilisateur, "
        "la barre de recherche globale et l'accès au centre de notifications.")

    add_heading(doc, "6.3  Composants réactifs (Alpine.js)", 2)
    add_bullet(doc, "Formulaire de projet : ajout dynamique de lignes de financement sans rechargement.")
    add_bullet(doc, "Moteur d'analyse du tableau de bord : filtres et sélection des axes en temps réel.")
    add_bullet(doc, "Interface d'import : prévisualisation accordion par section (Bailleurs, Projets, etc.).")
    add_bullet(doc, "Interface assistant IA : saisie et affichage du chat en temps réel.")
    add_bullet(doc, "Notifications : chargement asynchrone à l'ouverture du panneau.")
    add_bullet(doc, "Recherche globale : résultats en temps réel avec icônes différenciatrices par type de résultat.")

    add_heading(doc, "6.4  Messages utilisateur", 2)
    add_paragraph(doc,
        "Les confirmations de succès, avertissements et erreurs sont gérés par le système de messages "
        "Django (django.contrib.messages) et affichés sous forme de bannières colorées en haut de page, "
        "avec fermeture automatique ou manuelle.")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Déploiement
# ════════════════════════════════════════════════════════════════════════════

def build_section_7(doc):
    add_heading(doc, "7.  Déploiement et infrastructure", 1)
    add_separator(doc)

    add_heading(doc, "7.1  Environnements", 2)
    envs = [
        ("Développement", "SQLite local, DEBUG=True, runserver Django", "Machine du développeur"),
        ("Production", "PostgreSQL, DEBUG=False, Gunicorn + WhiteNoise", "Render.com (cloud PaaS)"),
    ]
    add_multi_col_table(doc, ["Environnement", "Configuration", "Infrastructure"], envs, col_widths=[3.5, 7.5, 5.5])

    add_heading(doc, "7.2  Variables d'environnement", 2)
    env_vars = [
        ("SECRET_KEY", "Clé secrète Django — OBLIGATOIRE en production."),
        ("DEBUG", "Valeur 'True' ou 'False' — 'False' en production."),
        ("DATABASE_URL", "URL complète de la base PostgreSQL (format dj-database-url). Si absente, SQLite est utilisé."),
        ("ALLOWED_HOSTS", "Géré via ALLOWED_HOSTS dans settings.py (*.onrender.com, *.ngrok-free.app)."),
        ("GEMINI_API_KEY", "Clé API Google Gemini pour l'assistant IA — facultative (désactive l'IA si absente)."),
    ]
    add_multi_col_table(doc, ["Variable", "Description"], env_vars, col_widths=[4.5, 12])

    add_heading(doc, "7.3  Procédure de déploiement", 2)
    add_paragraph(doc, "La procédure standard de déploiement suit les étapes suivantes :")
    steps = [
        "1. Cloner le dépôt Git sur la plateforme cible.",
        "2. Configurer les variables d'environnement (SECRET_KEY, DATABASE_URL, GEMINI_API_KEY, DEBUG=False).",
        "3. Installer les dépendances : pip install -r requirements.txt.",
        "4. Appliquer les migrations : python manage.py migrate.",
        "5. Collecter les fichiers statiques : python manage.py collectstatic --noinput.",
        "6. Créer le compte super administrateur : python manage.py createsuperuser.",
        "7. (Optionnel) Charger les données initiales via l'interface d'import Excel.",
        "8. Lancer Gunicorn : gunicorn config.wsgi:application --bind 0.0.0.0:$PORT.",
    ]
    for s in steps:
        add_bullet(doc, s)

    add_info_box(doc,
        "ℹ️  Le fichier render.yaml à la racine du projet contient la configuration déclarative pour "
        "le déploiement automatisé sur Render.com. Le fichier build.sh contient les commandes de build "
        "exécutées automatiquement lors de chaque déploiement (install, migrate, collectstatic).",
        bg='FFF7ED', border='F77F00')

    add_heading(doc, "7.4  Gestion des fichiers statiques", 2)
    add_paragraph(doc,
        "WhiteNoise est configuré comme middleware de service des fichiers statiques, permettant de servir "
        "les assets CSS, JavaScript et images directement depuis le processus Gunicorn, sans serveur web "
        "frontal séparé. La compression et le cache HTTP long terme sont activés en production "
        "(CompressedManifestStaticFilesStorage).")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — Exigences non fonctionnelles
# ════════════════════════════════════════════════════════════════════════════

def build_section_8(doc):
    add_heading(doc, "8.  Exigences non fonctionnelles", 1)
    add_separator(doc)

    exigences = [
        ("Performance",
         "Les pages principales (tableau de bord, liste des projets) doivent se charger en moins de 3 secondes "
         "pour un portefeuille de 500 projets. Les données analytiques sont pré-sérialisées en JSON côté serveur "
         "pour un rendu côté client sans requêtes supplémentaires."),
        ("Disponibilité",
         "L'application doit être disponible 24h/24, 7j/7. En cas d'indisponibilité de l'API Gemini (quota dépassé), "
         "l'assistant IA affiche un message d'erreur explicite sans impacter les autres fonctionnalités."),
        ("Compatibilité navigateurs",
         "Support des versions récentes de Google Chrome, Mozilla Firefox et Microsoft Edge. "
         "L'interface n'est pas optimisée pour les appareils mobiles dans sa version actuelle."),
        ("Sécurité des données",
         "Toutes les communications doivent se faire via HTTPS en production. "
         "La clé secrète Django et les credentials API ne doivent jamais être exposés dans le code source. "
         "Les domaines de confiance CSRF sont explicitement configurés."),
        ("Internationalisation",
         "L'interface est intégralement en langue française. Le format de date est DD/MM/YYYY. "
         "Le séparateur décimal en français (virgule) est géré côté serveur via json.dumps() pour "
         "éviter des erreurs de parsing JavaScript."),
        ("Maintenabilité",
         "Le code respecte les conventions Django (séparation models/views/urls/templates, ORM, "
         "generic patterns). Les calculs financiers critiques sont des propriétés Python (@property) "
         "centralisées dans les modèles."),
        ("Extensibilité",
         "L'architecture modulaire Django permet d'ajouter de nouveaux modules (ex. reporting PDF, "
         "notifications email, API REST) sans modifier le cœur applicatif."),
        ("Gestion des erreurs",
         "Les erreurs de validation de formulaire sont affichées en ligne. Les erreurs système "
         "(404, 403, 500) sont gérées par des pages Django standards ou personnalisables."),
    ]

    for titre, desc in exigences:
        add_heading(doc, titre, 3)
        add_paragraph(doc, desc)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — Contraintes et dépendances
# ════════════════════════════════════════════════════════════════════════════

def build_section_9(doc):
    add_heading(doc, "9.  Contraintes et dépendances", 1)
    add_separator(doc)

    add_heading(doc, "9.1  Dépendances externes", 2)
    deps = [
        ("API Google Gemini", "Nécessite une clé API valide. En cas d'absence ou d'expiration de la clé, l'assistant IA est non fonctionnel mais les autres modules restent opérationnels."),
        ("CDN TailwindCSS, Alpine.js, ApexCharts, Leaflet, Material Symbols", "L'interface dépend de ressources CDN. En environnement sans accès internet, ces ressources doivent être hébergées localement."),
        ("PostgreSQL (production)", "La migration depuis SQLite vers PostgreSQL est transparente grâce à dj-database-url. La compatibilité est assurée par le framework Django ORM."),
    ]
    add_multi_col_table(doc, ["Dépendance", "Contrainte"], deps, col_widths=[4.5, 12])

    add_heading(doc, "9.2  Contraintes métier", 2)
    add_bullet(doc, "Les taux de conversion des devises sont des valeurs indicatives codées en dur. Ils doivent être mis à jour périodiquement par un administrateur technique pour refléter les cours réels.")
    add_bullet(doc, "La détection des projets en retard est basée sur la date système du serveur (fuseau Africa/Ouagadougou). Les serveurs doivent être synchronisés sur ce fuseau.")
    add_bullet(doc, "Le plan actif du PND est déterminé par le flag actif=True. Un seul plan peut être actif à la fois dans l'application. La gestion multi-plans est prévue dans l'interface d'administration Django.")
    add_bullet(doc, "Les données de co-financement (plusieurs bailleurs pour un même projet) sont gérées via la table Financement, distinctement du champ bailleur_principal qui indique uniquement le bailleur chef de file.")

    add_heading(doc, "9.3  Limites actuelles et évolutions prévues", 2)
    add_bullet(doc, "Pas de génération de rapports PDF natifs (possible via WeasyPrint ou ReportLab en extension).")
    add_bullet(doc, "Pas de notifications par email (possible via Django email backend + SMTP).")
    add_bullet(doc, "Pas d'API REST externe consommable par des systèmes tiers (possible via Django REST Framework).")
    add_bullet(doc, "Interface non optimisée pour les écrans mobiles (responsive partiel).")
    add_bullet(doc, "L'assistant IA n'a pas de mémoire persistante entre les sessions (l'historique est en mémoire de session uniquement).")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — Glossaire
# ════════════════════════════════════════════════════════════════════════════

def build_section_10(doc):
    add_heading(doc, "10.  Glossaire", 1)
    add_separator(doc)

    terms = [
        ("APD", "Aide Publique au Développement — flux financiers publics en faveur du développement des pays."),
        ("Bailleur de fonds", "Institution (multilatérale, bilatérale, ONG, etc.) qui finance des projets de développement."),
        ("Co-financement", "Projet financé simultanément par plusieurs bailleurs, chacun avec un accord de financement distinct."),
        ("Décaissement", "Versement effectif d'une tranche de financement accordé à un projet."),
        ("Devise", "Monnaie dans laquelle est libellé un accord de financement (USD, EUR, XOF, etc.)."),
        ("Engagement", "Montant total d'un accord de financement signé entre un bailleur et le gouvernement."),
        ("Financement", "Accord entre un bailleur et un projet définissant le montant engagé, le type et les modalités."),
        ("Gemini", "Modèle de langage (LLM) développé par Google, utilisé pour l'assistant IA de l'application."),
        ("KPI", "Key Performance Indicator — Indicateur clé de performance."),
        ("Pilier", "Axe stratégique principal d'un Plan National de Développement."),
        ("PND", "Plan National de Développement — Document de planification stratégique national."),
        ("Point Focal", "Agent de l'administration assigné à un ou plusieurs bailleurs pour le suivi opérationnel."),
        ("PTF", "Partenaires Techniques et Financiers — ensemble des bailleurs de fonds."),
        ("RBAC", "Role-Based Access Control — Contrôle d'accès basé sur les rôles."),
        ("Reste à décaisser", "Différence entre le montant engagé et le montant déjà décaissé."),
        ("Secteur", "Domaine d'intervention d'un projet (Agriculture, Santé, Éducation, Infrastructure, etc.)."),
        ("Sous-objectif", "Objectif spécifique rattaché à un pilier du PND, utilisé pour l'alignement des projets."),
        ("Taux d'avancement physique", "Pourcentage de réalisation physique du projet (travaux, activités) renseigné manuellement."),
        ("Taux de décaissement", "Ratio en pourcentage du montant décaissé sur le montant total engagé."),
        ("XOF", "Franc CFA de l'Afrique de l'Ouest — monnaie nationale de la Côte d'Ivoire."),
    ]
    add_two_col_table(doc, terms, header=["Terme", "Définition"])


# ════════════════════════════════════════════════════════════════════════════
#  POINT D'ENTRÉE
# ════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    doc = build_document()
    output_path = 'Cahier_des_Charges_Suivi_Projets_Bailleurs.docx'
    doc.save(output_path)
    print(f"✅  Document généré : {output_path}")
