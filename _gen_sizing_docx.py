"""
Génère la Fiche de Sizing au format Word (.docx)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Couleurs ────────────────────────────────────────────────────────────────
CI_ORANGE   = RGBColor(0xF7, 0x7F, 0x00)
CI_GREEN    = RGBColor(0x00, 0x9A, 0x44)
CI_DARK     = RGBColor(0x1E, 0x29, 0x3B)
CI_LIGHT    = RGBColor(0xF8, 0xFA, 0xFC)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_HEADER = RGBColor(0xE2, 0xE8, 0xF0)
GRAY_ROW    = RGBColor(0xF8, 0xFA, 0xFC)
ORANGE_SOFT = RGBColor(0xFF, 0xED, 0xD5)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CBD5E1')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, italic=False, size=10,
              color: RGBColor = CI_DARK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color


def add_styled_table(doc, headers, rows, col_widths=None):
    """Crée un tableau stylisé avec en-tête coloré orange et lignes alternées."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # En-tête
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], CI_ORANGE)
        set_cell_borders(hdr_cells[i])
        cell_text(hdr_cells[i], h, bold=True, color=WHITE, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Lignes de données
    for r_idx, row in enumerate(rows):
        bg = GRAY_ROW if r_idx % 2 == 0 else WHITE
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            set_cell_bg(row_cells[c_idx], bg)
            set_cell_borders(row_cells[c_idx])
            is_first = (c_idx == 0)
            cell_text(row_cells[c_idx], str(val), bold=is_first, size=10,
                      color=CI_DARK)

    # Largeurs colonnes
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Cm(w)

    return table


def add_section_title(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run_num = p.add_run(f"{number}. ")
    run_num.bold = True
    run_num.font.size  = Pt(13)
    run_num.font.color.rgb = CI_ORANGE
    run_title = p.add_run(title.upper())
    run_title.bold = True
    run_title.font.size  = Pt(13)
    run_title.font.color.rgb = CI_DARK


def add_subsection_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size  = Pt(11)
    run.font.color.rgb = CI_GREEN


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run("ℹ  " + text)
    run.italic = True
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def add_warning(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run("⚠  " + text)
    run.italic = True
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)


def add_error_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run("✖  " + text)
    run.italic = True
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)


# ═══════════════════════════════════════════════════════════════════════════
# CONSTRUCTION DU DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Marges ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Bandeau d'en-tête (tableau 1 ligne) ─────────────────────────────────────
header_tbl = doc.add_table(rows=2, cols=1)
header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

cell_top = header_tbl.rows[0].cells[0]
set_cell_bg(cell_top, CI_DARK)
p = cell_top.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("FICHE DE SIZING — HÉBERGEMENT SERVEUR")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = WHITE

cell_sub = header_tbl.rows[1].cells[0]
set_cell_bg(cell_sub, CI_ORANGE)
p2 = cell_sub.paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(4)
p2.paragraph_format.space_after  = Pt(4)
r2 = p2.add_run("Plateforme de Suivi des Projets des Bailleurs de Fonds — Mai 2026")
r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = WHITE

doc.add_paragraph()  # espace

# ── Bloc destinataire ────────────────────────────────────────────────────────
meta_tbl = doc.add_table(rows=2, cols=2)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
meta_data = [
    ("Destinataire :", "Service Informatique du Ministère"),
    ("Date :",         "Mai 2026"),
]
for i, (label, val) in enumerate(meta_data):
    set_cell_bg(meta_tbl.rows[i].cells[0], ORANGE_SOFT)
    cell_text(meta_tbl.rows[i].cells[0], label, bold=True, size=10, color=CI_DARK)
    cell_text(meta_tbl.rows[i].cells[1], val, size=10, color=CI_DARK)
    for c in meta_tbl.rows[i].cells:
        set_cell_borders(c)
for r in meta_tbl.rows:
    r.cells[0].width = Cm(5)
    r.cells[1].width = Cm(12)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1 — RESSOURCES MATÉRIELLES
# ═══════════════════════════════════════════════════════════════════════════

add_section_title(doc, "1", "Ressources matérielles")

# 1.1 Serveur application
add_subsection_title(doc, "1.1  Serveur d'application")

add_styled_table(doc,
    headers=["Ressource", "Minimum requis", "Recommandé (production)"],
    rows=[
        ["vCPU",                         "2 vCPU",        "4 vCPU"],
        ["RAM",                          "2 Go",          "4 Go"],
        ["Stockage OS + application",    "10 Go SSD",     "20 Go SSD"],
        ["Stockage données (media/)",    "20 Go HDD/SSD", "50 Go SSD"],
        ["Bande passante réseau",        "10 Mbps",       "100 Mbps"],
    ],
    col_widths=[7, 4.5, 5.5]
)

add_note(doc,
    "Justification CPU/RAM : Gunicorn (serveur WSGI Python) est dimensionné selon la formule "
    "2 × nb_vCPU + 1 workers. Avec 4 vCPU → 9 workers, chaque worker consomme ~150–200 Mo de RAM. "
    "Pour 10 à 50 utilisateurs simultanés (usage interne ministériel), 4 Go sont largement suffisants."
)
add_note(doc,
    "Justification stockage : Le répertoire /media/ accueille les logos des bailleurs et les "
    "pièces jointes des projets (PDF, DOCX, XLSX). 50 Go offre une marge confortable pour plusieurs "
    "années d'exploitation."
)

doc.add_paragraph()

# 1.2 Serveur BDD
add_subsection_title(doc, "1.2  Serveur de base de données")
p = doc.add_paragraph()
r = p.add_run("Peut être mutualisé sur le même serveur physique si les ressources le permettent.")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

add_styled_table(doc,
    headers=["Ressource", "Minimum requis", "Recommandé (production)"],
    rows=[
        ["vCPU",                   "2 vCPU",     "4 vCPU"],
        ["RAM",                    "1 Go",        "4 Go"],
        ["Stockage base de données","5 Go SSD",   "20 Go SSD"],
        ["Réseau",                 "LAN interne", "LAN interne"],
    ],
    col_widths=[7, 4.5, 5.5]
)

add_note(doc,
    "Justification : La base héberge des tables structurées de taille modérée (projets, financements, "
    "décaissements, utilisateurs, audit log). Pour plusieurs centaines de projets, le volume de données "
    "reste inférieur à 1 Go. 20 Go couvrent plusieurs années avec une grande marge."
)

doc.add_paragraph()

# 1.3 Récapitulatif serveur unique
add_subsection_title(doc, "1.3  Configuration cible — Serveur unique mutualisé")
p = doc.add_paragraph()
r = p.add_run("Si un seul serveur doit tout héberger (application + base de données) :")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = CI_DARK

add_styled_table(doc,
    headers=["Ressource", "Configuration cible"],
    rows=[
        ["vCPU",          "4 vCPU"],
        ["RAM",           "8 Go"],
        ["Stockage total", "80 Go SSD  (OS 20 Go + App 10 Go + BDD 10 Go + Media 40 Go)"],
        ["Bande passante", "100 Mbps"],
        ["Type de disque", "SSD (obligatoire pour la base de données)"],
    ],
    col_widths=[6, 11]
)

doc.add_paragraph()

# Saut de page avant section 2
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2 — RESSOURCES LOGICIELLES
# ═══════════════════════════════════════════════════════════════════════════

add_section_title(doc, "2", "Ressources logicielles")

# 2.1 OS
add_subsection_title(doc, "2.1  Système d'exploitation")
add_styled_table(doc,
    headers=["Paramètre", "Valeur retenue"],
    rows=[
        ["OS recommandé",    "Ubuntu Server 22.04 LTS (Jammy Jellyfish)"],
        ["Alternative",      "Debian 12 (Bookworm)"],
        ["Architecture",     "x86_64 (64 bits)"],
        ["Mode d'installation", "Serveur (sans interface graphique)"],
    ],
    col_widths=[6, 11]
)
doc.add_paragraph()

# 2.2 Runtimes
add_subsection_title(doc, "2.2  Runtimes applicatifs")
add_styled_table(doc,
    headers=["Runtime", "Version requise", "Notes"],
    rows=[
        ["Python",    "3.11.x",      "Version exacte requise par l'application"],
        ["pip",       "≥ 23.x",      "Gestionnaire de paquets Python"],
        ["venv",      "Natif Python 3.11", "Isolation de l'environnement"],
    ],
    col_widths=[4, 4, 9]
)
add_warning(doc,
    "Python 3.11 est impératif. Python 3.10 et inférieurs ne sont pas supportés. "
    "Les versions 3.12+ nécessitent une vérification de compatibilité préalable."
)
doc.add_paragraph()

# 2.3 BDD
add_subsection_title(doc, "2.3  Moteur de base de données")
add_styled_table(doc,
    headers=["Paramètre", "Valeur retenue"],
    rows=[
        ["SGBD",                  "PostgreSQL"],
        ["Version minimale",      "PostgreSQL 14"],
        ["Version recommandée",   "PostgreSQL 15 ou 16"],
        ["Pilote Python",         "psycopg2-binary ≥ 2.9.9  (installé automatiquement via pip)"],
        ["Dépendance système",    "libpq-dev  (bibliothèque client PostgreSQL)"],
        ["Encodage",              "UTF-8"],
        ["Nom de la base",        "appli_suivi_projets  (configurable)"],
    ],
    col_widths=[6, 11]
)
add_error_note(doc,
    "MySQL / MariaDB non supportés — l'application est conçue exclusivement pour PostgreSQL en production."
)
add_note(doc,
    "SQLite est disponible uniquement pour les tests locaux, jamais pour la production."
)
doc.add_paragraph()

# 2.4 Serveur web
add_subsection_title(doc, "2.4  Serveur web / Reverse proxy")
add_styled_table(doc,
    headers=["Composant", "Rôle", "Version"],
    rows=[
        ["Nginx",    "Reverse proxy, SSL/TLS, fichiers statiques",     "≥ 1.18"],
        ["Gunicorn", "Serveur WSGI Python (Django ↔ Nginx)",           "≥ 21.2.0"],
    ],
    col_widths=[4, 9, 4]
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run("Flux de requête : ")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = CI_DARK
r2 = p.add_run("Navigateur  →  Nginx (HTTPS:443)  →  Gunicorn (HTTP:8000)  →  Django  →  PostgreSQL")
r2.font.size = Pt(10); r2.font.color.rgb = CI_DARK; r2.bold = False

add_note(doc,
    "Gunicorn est installé automatiquement via pip (requirements.txt). "
    "Nginx doit être installé séparément via le gestionnaire de paquets système (apt)."
)
doc.add_paragraph()

# 2.5 Dépendances système
add_subsection_title(doc, "2.5  Dépendances système Linux (packages apt)")
add_styled_table(doc,
    headers=["Package", "Usage"],
    rows=[
        ["python3.11  +  python3.11-venv  +  python3.11-dev", "Langage principal de l'application"],
        ["python3-pip",          "Gestionnaire de paquets Python"],
        ["libpq-dev",            "Bibliothèque cliente PostgreSQL (requis par psycopg2)"],
        ["gcc  +  build-essential", "Compilation des extensions Python"],
        ["nginx",                "Serveur web / reverse proxy"],
        ["postgresql  +  postgresql-contrib", "Moteur de base de données"],
        ["git",                  "Déploiement et mises à jour du code source (optionnel)"],
        ["curl",                 "Utilitaire réseau (optionnel)"],
    ],
    col_widths=[8, 9]
)
doc.add_paragraph()

# 2.6 Dépendances Python
add_subsection_title(doc, "2.6  Dépendances Python (fichier requirements.txt)")
p = doc.add_paragraph()
r = p.add_run("Toutes les dépendances Python s'installent avec une seule commande :  ")
r.font.size = Pt(10); r.font.color.rgb = CI_DARK
r2 = p.add_run("pip install -r requirements.txt")
r2.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = CI_GREEN

add_styled_table(doc,
    headers=["Paquet Python", "Version", "Usage"],
    rows=[
        ["django",               "≥ 5.2, < 6.0", "Framework web principal"],
        ["gunicorn",             "≥ 21.2.0",     "Serveur WSGI de production"],
        ["psycopg2-binary",      "≥ 2.9.9",      "Connecteur PostgreSQL"],
        ["openpyxl",             "≥ 3.1.0",      "Lecture/écriture fichiers Excel (.xlsx)"],
        ["whitenoise",           "≥ 6.6.0",      "Service des fichiers statiques"],
        ["dj-database-url",      "≥ 2.1.0",      "Parsing URL de connexion base de données"],
        ["google-generativeai",  "≥ 0.3.0",      "Client API Google Gemini (Assistant IA — optionnel)"],
    ],
    col_widths=[5, 3.5, 8.5]
)
doc.add_paragraph()

# 2.7 Ports réseau
add_subsection_title(doc, "2.7  Protocoles et ports réseau requis")
add_styled_table(doc,
    headers=["Port", "Protocole", "Direction", "Usage"],
    rows=[
        ["443",  "HTTPS/TLS", "Entrant (internet → serveur)", "Accès navigateurs (production)"],
        ["80",   "HTTP",      "Entrant (internet → serveur)", "Redirection automatique vers HTTPS"],
        ["5432", "TCP",       "Interne serveur",              "Communication Django ↔ PostgreSQL"],
        ["8000", "HTTP",      "Interne serveur",              "Communication Nginx ↔ Gunicorn"],
    ],
    col_widths=[2, 3, 6, 6]
)
doc.add_paragraph()

# 2.8 Certificat SSL
add_subsection_title(doc, "2.8  Certificat SSL/TLS")
p = doc.add_paragraph()
r = p.add_run("Un certificat SSL/TLS est obligatoire pour la mise en production.")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)

add_styled_table(doc,
    headers=["Option", "Description"],
    rows=[
        ["Certificat PKI interne du Ministère",
         "Si le Ministère dispose d'une autorité de certification (solution préférée)"],
        ["Let's Encrypt (Certbot)",
         "Certificat gratuit et automatique — nécessite un accès internet sortant"],
        ["Certificat commercial",
         "Achat auprès d'un tiers de confiance (DigiCert, Sectigo, etc.)"],
    ],
    col_widths=[6, 11]
)
doc.add_paragraph()

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3 — RÉCAPITULATIF GLOBAL
# ═══════════════════════════════════════════════════════════════════════════

add_section_title(doc, "3", "Récapitulatif global")
doc.add_paragraph()

summary_tbl = doc.add_table(rows=14, cols=2)
summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

summary_data = [
    # (label, value, is_section_header)
    ("MATÉRIEL",                  "",                                          True),
    ("vCPU",                      "4 vCPU  (minimum : 2)",                    False),
    ("RAM",                       "8 Go  (minimum : 4 Go)",                   False),
    ("Stockage",                  "80 Go SSD  (minimum : 40 Go)",             False),
    ("Bande passante",            "100 Mbps",                                 False),
    ("LOGICIEL",                  "",                                          True),
    ("Système d'exploitation",    "Ubuntu Server 22.04 LTS  (x86_64)",        False),
    ("Langage",                   "Python 3.11.x",                            False),
    ("Framework",                 "Django 5.2  (installé via pip)",           False),
    ("Serveur WSGI",              "Gunicorn 21.2+  (installé via pip)",       False),
    ("Reverse proxy",             "Nginx ≥ 1.18",                             False),
    ("Base de données",           "PostgreSQL 15 ou 16",                      False),
    ("CONNECTIVITÉ",              "",                                          True),
    ("Ports / SSL",               "80 (HTTP), 443 (HTTPS), certificat SSL/TLS requis", False),
]

for i, (label, value, is_header) in enumerate(summary_data):
    cells = summary_tbl.rows[i].cells
    if is_header:
        merged = cells[0].merge(cells[1])
        set_cell_bg(merged, CI_DARK)
        cell_text(merged, "  " + label, bold=True, size=11,
                  color=WHITE, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_borders(merged)
    else:
        bg = GRAY_ROW if i % 2 == 0 else WHITE
        set_cell_bg(cells[0], ORANGE_SOFT)
        set_cell_bg(cells[1], bg)
        cell_text(cells[0], label, bold=True, size=10, color=CI_DARK)
        cell_text(cells[1], value, size=10, color=CI_DARK)
        set_cell_borders(cells[0])
        set_cell_borders(cells[1])

for r in summary_tbl.rows:
    r.cells[0].width = Cm(6)
    try:
        r.cells[1].width = Cm(11)
    except Exception:
        pass

doc.add_paragraph()

# ── Pied de document ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Document établi sur la base de l'analyse du code source de l'application — Mai 2026\n"
    "Plateforme de Suivi des Projets des Bailleurs de Fonds — Cabinet du Ministère en charge de la Planification"
)
run.italic = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ── Sauvegarde ───────────────────────────────────────────────────────────────
output_path = "FICHE_SIZING_HEBERGEMENT.docx"
doc.save(output_path)
print(f"✅ Fichier généré : {output_path}")
