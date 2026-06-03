"""
Génération du Cahier des Charges — Plateforme de Suivi des Projets Bailleurs
Ministère du Plan et du Développement — République de Côte d'Ivoire
Version 2.0 — Document destiné à la consultation/commande d'un prestataire
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Palette de couleurs ────────────────────────────────────────────────────────
VERT_CI       = RGBColor(0x00, 0x9A, 0x44)   # Vert (drapeau CI)
ORANGE_CI     = RGBColor(0xF7, 0x7F, 0x00)   # Orange accent
DARK_BLUE     = RGBColor(0x0D, 0x2E, 0x4E)   # Bleu marine profond
MID_BLUE      = RGBColor(0x1A, 0x5E, 0x8A)   # Bleu institutionnel
LIGHT_BLUE    = RGBColor(0xE8, 0xF4, 0xFB)   # Fond bleu très clair
GOLD          = RGBColor(0xD4, 0xA0, 0x17)   # Or accent
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BLACK         = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_DARK     = RGBColor(0x4A, 0x4A, 0x4A)
GRAY_MID      = RGBColor(0x7A, 0x8A, 0x9A)
BG_ORANGE     = RGBColor(0xFF, 0xF3, 0xE0)   # Fond orange très clair
BG_GREEN      = RGBColor(0xE8, 0xF5, 0xE9)   # Fond vert très clair


# ══════════════════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='C5D3E0', size='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def no_borders(table):
    """Supprime toutes les bordures d'un tableau."""
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_run(para, text, bold=False, italic=False, size=10, color=None, underline=False):
    r = para.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    r.font.underline = underline
    r.font.size = Pt(size)
    r.font.color.rgb = color if color else BLACK
    r.font.name = 'Calibri'
    return r

def styled_para(doc, text='', bold=False, italic=False, size=10.5, color=None,
                align=WD_ALIGN_PARAGRAPH.LEFT, space_before=3, space_after=6,
                indent_left=0, indent_first=0, line_spacing=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_left:
        pf.left_indent = Cm(indent_left)
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    if line_spacing:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p

def heading1(doc, text, num=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    # Barre latérale orange
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), 'F77F00')
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent = Cm(0.5)
    if num:
        r1 = p.add_run(f'{num}  ')
        r1.font.bold = True
        r1.font.size = Pt(15)
        r1.font.color.rgb = ORANGE_CI
        r1.font.name = 'Calibri'
    r2 = p.add_run(text.upper())
    r2.font.bold = True
    r2.font.size = Pt(15)
    r2.font.color.rgb = DARK_BLUE
    r2.font.name = 'Calibri'
    return p

def heading2(doc, text, num=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    if num:
        r1 = p.add_run(f'{num}  ')
        r1.font.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = ORANGE_CI
        r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = MID_BLUE
    r2.font.name = 'Calibri'
    return p

def heading3(doc, text, num=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    if num:
        r1 = p.add_run(f'{num}  ')
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = DARK_BLUE
        r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.bold = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK_BLUE
    r2.font.name = 'Calibri'
    return p

def body(doc, text, size=10.5, space_after=8, indent=0, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = GRAY_DARK
    r.font.name = 'Calibri'
    return p

def bullet(doc, text, bold_prefix=None, level=0, size=10.5, indent=1.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(indent + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    bullet_char = '▸' if level == 0 else '–'
    r0 = p.add_run(f'{bullet_char}  ')
    r0.font.size = Pt(size)
    r0.font.color.rgb = ORANGE_CI if level == 0 else MID_BLUE
    r0.font.name = 'Calibri'
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(size)
        r1.font.color.rgb = DARK_BLUE
        r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    r2.font.color.rgb = GRAY_DARK
    r2.font.name = 'Calibri'
    return p

def numbered_item(doc, num, text, size=10.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    r1 = p.add_run(f'{num}.  ')
    r1.font.bold = True
    r1.font.size = Pt(size)
    r1.font.color.rgb = ORANGE_CI
    r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    r2.font.color.rgb = GRAY_DARK
    r2.font.name = 'Calibri'
    return p

def info_box(doc, title, text, bg='E8F4FB', border_color='1A5E8A', title_color=None):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0, 0)
    set_cell_bg(c, bg)
    set_cell_borders(c, border_color, '8')
    c.paragraphs[0].clear()
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.4)
    if title:
        r1 = p.add_run(f'{title}  ')
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = title_color if title_color else MID_BLUE
        r1.font.name = 'Calibri'
    # body paragraphs
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i == 0 and not title:
            pr = p
        elif i == 0:
            pr = c.add_paragraph()
        else:
            pr = c.add_paragraph()
        pr.paragraph_format.left_indent = Cm(0.4)
        pr.paragraph_format.space_before = Pt(1)
        pr.paragraph_format.space_after = Pt(4 if i == len(lines)-1 else 1)
        if i > 0 or title:
            r = pr.add_run(line)
            r.font.size = Pt(9.5)
            r.font.color.rgb = GRAY_DARK
            r.font.name = 'Calibri'
        else:
            r = p.add_run(line)
            r.font.size = Pt(9.5)
            r.font.color.rgb = GRAY_DARK
            r.font.name = 'Calibri'
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return t

def separator(doc, color='F77F00'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def big_table(doc, headers, rows, col_widths=None, header_bg='0D2E4E', alt_bg='EFF4F9', font_size=9.5):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell, header_bg, '2')
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = WHITE
        r.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for i, row in enumerate(rows):
        bg = alt_bg if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row):
            cell = t.cell(i+1, j)
            set_cell_bg(cell, bg)
            set_cell_borders(cell, 'C5D3E0', '4')
            p = cell.paragraphs[0]
            p.clear()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.15)
            if isinstance(val, tuple):
                r = p.add_run(val[0])
                r.font.bold = val[1]
            else:
                r = p.add_run(str(val))
                r.font.bold = (j == 0 and len(headers) > 1)
            r.font.size = Pt(font_size)
            r.font.color.rgb = GRAY_DARK if j > 0 else DARK_BLUE
            r.font.name = 'Calibri'

    if col_widths:
        for row in t.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return t

def profile_card(doc, title, subtitle, items, bg_hex='EFF4F9', accent='1A5E8A'):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0, 0)
    set_cell_bg(c, bg_hex)
    set_cell_borders(c, accent, '10')
    c.paragraphs[0].clear()
    # Title
    pt = c.paragraphs[0]
    pt.paragraph_format.space_before = Pt(6)
    pt.paragraph_format.space_after = Pt(2)
    pt.paragraph_format.left_indent = Cm(0.4)
    r1 = pt.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(11.5)
    r1.font.color.rgb = RGBColor.from_string(accent) if isinstance(accent, str) else MID_BLUE
    r1.font.name = 'Calibri'
    # Subtitle
    ps = c.add_paragraph()
    ps.paragraph_format.left_indent = Cm(0.4)
    ps.paragraph_format.space_before = Pt(0)
    ps.paragraph_format.space_after = Pt(4)
    r2 = ps.add_run(subtitle)
    r2.font.italic = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GRAY_MID
    r2.font.name = 'Calibri'
    # Items
    for item in items:
        pi = c.add_paragraph()
        pi.paragraph_format.left_indent = Cm(0.8)
        pi.paragraph_format.first_line_indent = Cm(-0.4)
        pi.paragraph_format.space_before = Pt(1)
        pi.paragraph_format.space_after = Pt(3)
        ri = pi.add_run('▸  ')
        ri.font.size = Pt(9.5)
        ri.font.color.rgb = ORANGE_CI
        ri.font.name = 'Calibri'
        ri2 = pi.add_run(item)
        ri2.font.size = Pt(9.5)
        ri2.font.color.rgb = GRAY_DARK
        ri2.font.name = 'Calibri'
    # Bottom padding
    pe = c.add_paragraph()
    pe.paragraph_format.space_after = Pt(4)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return t


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════

def cover_page(doc):
    # Bandeau tricolore CI
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_borders(t)
    colors = ['F77F00', 'FFFFFF', '009A44']
    for j, col in enumerate(colors):
        c = t.cell(0, j)
        set_cell_bg(c, col)
        c.paragraphs[0].paragraph_format.space_before = Pt(5)
        c.paragraphs[0].paragraph_format.space_after = Pt(5)
        c.width = Cm(5.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Entête institutionnelle
    t2 = doc.add_table(rows=1, cols=1)
    t2.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_borders(t2)
    c2 = t2.cell(0, 0)
    set_cell_bg(c2, '0D2E4E')
    c2.width = Cm(16.5)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(12)
    p2.paragraph_format.space_after = Pt(4)
    r = p2.add_run("RÉPUBLIQUE DE CÔTE D'IVOIRE")
    r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = WHITE; r.font.name = 'Calibri'
    p2b = c2.add_paragraph()
    p2b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2b.paragraph_format.space_before = Pt(0)
    p2b.paragraph_format.space_after = Pt(4)
    rb = p2b.add_run("Union – Discipline – Travail")
    rb.font.italic = True; rb.font.size = Pt(10); rb.font.color.rgb = RGBColor(0xBF, 0xD7, 0xEA); rb.font.name = 'Calibri'
    p2c = c2.add_paragraph()
    p2c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2c.paragraph_format.space_before = Pt(4)
    p2c.paragraph_format.space_after = Pt(12)
    rc = p2c.add_run("Ministère du Plan et du Développement")
    rc.font.bold = True; rc.font.size = Pt(12); rc.font.color.rgb = RGBColor(0xD4, 0xE9, 0xF7); rc.font.name = 'Calibri'

    for _ in range(3): doc.add_paragraph()

    # Titre principal
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(0)
    pt.paragraph_format.space_after = Pt(8)
    rt = pt.add_run("CAHIER DES CHARGES")
    rt.font.bold = True; rt.font.size = Pt(28); rt.font.color.rgb = DARK_BLUE; rt.font.name = 'Calibri'

    # Ligne décorative orange
    tl = doc.add_table(rows=1, cols=1)
    tl.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(tl)
    cl = tl.cell(0, 0)
    set_cell_bg(cl, 'F77F00')
    cl.paragraphs[0].paragraph_format.space_before = Pt(3)
    cl.paragraphs[0].paragraph_format.space_after = Pt(3)
    cl.width = Cm(11)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Sous-titre
    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_before = Pt(4)
    ps.paragraph_format.space_after = Pt(4)
    rs = ps.add_run("Plateforme Numérique de Suivi des Projets et des Financements\ndes Partenaires Techniques et Financiers")
    rs.font.bold = True; rs.font.size = Pt(17); rs.font.color.rgb = MID_BLUE; rs.font.name = 'Calibri'

    ps2 = doc.add_paragraph()
    ps2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps2.paragraph_format.space_before = Pt(4)
    ps2.paragraph_format.space_after = Pt(2)
    rs2 = ps2.add_run("— Cabinet du Ministre —")
    rs2.font.italic = True; rs2.font.size = Pt(12); rs2.font.color.rgb = GRAY_MID; rs2.font.name = 'Calibri'

    for _ in range(3): doc.add_paragraph()

    # Bloc de métadonnées
    meta_data = [
        ("Commanditaire",  "Cabinet du Ministère du Plan et du Développement"),
        ("Objet du document", "Spécifications fonctionnelles et techniques — Phase 1"),
        ("Référence",      "MPD/CAB/CDC/2025-01"),
        ("Version",        "2.0 — Définitive"),
        ("Date",           datetime.date.today().strftime("%d %B %Y")),
        ("Statut",         "Document soumis à validation"),
        ("Destinataires",  "Prestataire retenu pour la réalisation — Directions partenaires"),
    ]
    tm = doc.add_table(rows=len(meta_data), cols=2)
    tm.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(meta_data):
        bg_l = '0D2E4E' if i % 2 == 0 else '1A5E8A'
        bg_v = 'EFF4F9' if i % 2 == 0 else 'FFFFFF'
        cl = tm.cell(i, 0); cv = tm.cell(i, 1)
        set_cell_bg(cl, bg_l); set_cell_borders(cl, bg_l, '2')
        set_cell_bg(cv, bg_v); set_cell_borders(cv, 'C5D3E0', '4')
        cl.width = Cm(5); cv.width = Cm(11.5)
        pl = cl.paragraphs[0]; pl.clear()
        pl.paragraph_format.left_indent = Cm(0.3)
        pl.paragraph_format.space_before = Pt(5); pl.paragraph_format.space_after = Pt(5)
        pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rl = pl.add_run(label)
        rl.font.bold = True; rl.font.size = Pt(9.5); rl.font.color.rgb = WHITE; rl.font.name = 'Calibri'
        pv = cv.paragraphs[0]; pv.clear()
        pv.paragraph_format.left_indent = Cm(0.3)
        pv.paragraph_format.space_before = Pt(5); pv.paragraph_format.space_after = Pt(5)
        rv = pv.add_run(val)
        rv.font.size = Pt(9.5); rv.font.color.rgb = DARK_BLUE; rv.font.name = 'Calibri'

    for _ in range(4): doc.add_paragraph()

    # Pied de garde
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("Cabinet du Ministre du Plan et du Développement  ·  Côte d'Ivoire  ·  2025\nDocument confidentiel — Usage interne")
    rf.font.italic = True; rf.font.size = Pt(8.5); rf.font.color.rgb = GRAY_MID; rf.font.name = 'Calibri'


# ══════════════════════════════════════════════════════════════════════════════
#  TABLE DES MATIÈRES
# ══════════════════════════════════════════════════════════════════════════════

def table_des_matieres(doc):
    heading1(doc, "Table des matières")
    separator(doc)
    doc.add_paragraph()

    toc = [
        # (numéro, texte, niveau 1/2/3)
        ("", "Avant-propos", 1),
        ("", "Liste des abréviations et acronymes", 1),
        ("1.", "Contexte général et enjeux institutionnels", 1),
        ("1.1", "Le Ministère du Plan et du Développement", 2),
        ("1.2", "Le rôle de gouverneur du Ministre et son portefeuille de bailleurs", 2),
        ("1.3", "Les principaux bailleurs concernés par la Phase 1", 2),
        ("1.4", "Diagnostic du dispositif actuel de suivi", 2),
        ("1.5", "Démarche consultative et parties prenantes", 2),
        ("2.", "Objectifs et vision stratégique de la plateforme", 1),
        ("2.1", "Vision et finalité de la plateforme", 2),
        ("2.2", "Objectifs opérationnels", 2),
        ("2.3", "Principes directeurs de conception", 2),
        ("3.", "Périmètre fonctionnel de la Phase 1", 1),
        ("3.1", "Définition du périmètre", 2),
        ("3.2", "Ce qui est hors périmètre (Phase 2 et au-delà)", 2),
        ("4.", "Profils utilisateurs et interfaces dédiées", 1),
        ("4.1", "Vue d'ensemble du système de gestion des accès", 2),
        ("4.2", "Interface Ministre — Vue synthétique de décision", 2),
        ("4.3", "Interface Directeur de Cabinet — Vue exhaustive de pilotage", 2),
        ("4.4", "Interface Point Focal / Cellule technique — Vue opérationnelle", 2),
        ("4.5", "Interface Conseiller / Lecteur — Vue de consultation", 2),
        ("5.", "Spécifications fonctionnelles détaillées", 1),
        ("5.1", "Tableau de bord et indicateurs de performance", 2),
        ("5.2", "Gestion des projets et programmes", 2),
        ("5.3", "Gestion des bailleurs de fonds", 2),
        ("5.4", "Financements et décaissements", 2),
        ("5.5", "Système d'alertes et de signalement", 2),
        ("5.6", "Module cartographique", 2),
        ("5.7", "Reporting et export des données", 2),
        ("5.8", "Import de données en masse", 2),
        ("5.9", "Assistant d'analyse par intelligence artificielle", 2),
        ("5.10", "Administration des comptes et des droits", 2),
        ("6.", "Exigences non fonctionnelles", 1),
        ("6.1", "Performance et temps de réponse", 2),
        ("6.2", "Sécurité des données et conformité réglementaire", 2),
        ("6.3", "Ergonomie et accessibilité", 2),
        ("6.4", "Maintenabilité et évolutivité", 2),
        ("6.5", "Disponibilité et continuité de service", 2),
        ("7.", "Architecture technique recommandée", 1),
        ("8.", "Modalités d'exécution et livrables attendus", 1),
        ("8.1", "Livrables", 2),
        ("8.2", "Planning indicatif", 2),
        ("8.3", "Profil du prestataire", 2),
        ("", "ANNEXE A — Recensement détaillé des bailleurs PTF (Phase 1)", 1),
        ("", "ANNEXE B — Synthèse des observations de la démarche consultative", 1),
        ("", "ANNEXE C — Glossaire des termes et abréviations", 1),
    ]
    for num, title, lvl in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0) if lvl == 1 else Cm(0.8)
        if num:
            r1 = p.add_run(f"{num}  ")
            r1.font.bold = (lvl == 1)
            r1.font.size = Pt(10.5 if lvl == 1 else 10)
            r1.font.color.rgb = ORANGE_CI if lvl == 1 else MID_BLUE
            r1.font.name = 'Calibri'
        r2 = p.add_run(title)
        r2.font.bold = (lvl == 1)
        r2.font.size = Pt(10.5 if lvl == 1 else 10)
        r2.font.color.rgb = DARK_BLUE
        r2.font.name = 'Calibri'


# ══════════════════════════════════════════════════════════════════════════════
#  AVANT-PROPOS
# ══════════════════════════════════════════════════════════════════════════════

def avant_propos(doc):
    heading1(doc, "Avant-propos")
    separator(doc)
    body(doc,
        "Le présent Cahier des Charges définit l'ensemble des spécifications fonctionnelles et techniques "
        "de la Plateforme Numérique de Suivi des Projets et des Financements des Partenaires Techniques et Financiers "
        "(PTF), dont la conception et le développement sont commandités par le Cabinet du Ministère du Plan et du Développement "
        "de la République de Côte d'Ivoire. Ce document constitue la référence contractuelle à partir de laquelle "
        "le prestataire retenu devra concevoir, développer, tester et livrer la solution.")
    body(doc,
        "Ce cahier des charges a été élaboré à l'issue d'une démarche consultative conduite du 5 au 8 mai 2026 "
        "auprès de l'ensemble des structures partenaires du Cabinet : Direction des Systèmes Informatiques et de la "
        "Digitalisation (DSID), Cellule de Coordination des Bailleurs Arabes, Cellule de Coordination et de Suivi du "
        "Portefeuille de la BAD (CCSPPP-BAD), Direction Générale du Plan (DGP), Direction Générale de la Coopération "
        "au Développement (DGCOD), Bureau National de la Prospective (BNPVS), ENSEA, et d'autres directions. "
        "L'ensemble des observations formulées lors de ces consultations a été intégré dans le présent document.")
    body(doc,
        "La plateforme a vocation à devenir le principal outil de pilotage stratégique et de veille active "
        "du portefeuille des projets de développement financés par les bailleurs de fonds, à destination "
        "du Ministre, du Directeur de Cabinet et des équipes techniques du Ministère. Elle sera conçue "
        "comme un outil évolutif, dont les fonctionnalités pourront être enrichies à l'issue de la Phase 1.")
    info_box(doc,
        "⚠  Note importante :",
        "Ce document est destiné à être partagé avec l'ensemble des parties prenantes pour amendement "
        "et validation avant transmission au prestataire. Toute observation complémentaire devra être "
        "adressée au Cabinet du Ministère du Plan et du Développement dans un délai de quinze (15) jours "
        "ouvrables à compter de la date de diffusion du présent document.",
        bg='FFF3E0', border_color='F77F00', title_color=ORANGE_CI)


# ══════════════════════════════════════════════════════════════════════════════
#  LISTE DES ABRÉVIATIONS
# ══════════════════════════════════════════════════════════════════════════════

def abreviations(doc):
    heading1(doc, "Liste des abréviations et acronymes")
    separator(doc)
    abbrevs = [
        ("ADB / BAD", "Banque Africaine de Développement / African Development Bank"),
        ("AFD",       "Agence Française de Développement"),
        ("AID / IDA", "Association Internationale de Développement (guichet BM)"),
        ("ANStat",    "Agence Nationale de la Statistique"),
        ("APD",       "Aide Publique au Développement"),
        ("BIRD",      "Banque Internationale pour la Reconstruction et le Développement (Banque mondiale)"),
        ("BIDC",      "Banque d'Investissement et de Développement de la CEDEAO"),
        ("BNPVS",     "Bureau National de la Prospective et de la Veille Stratégique"),
        ("BOAD",      "Banque Ouest Africaine de Développement"),
        ("CAD-OCDE",  "Comité d'Aide au Développement — Organisation de Coopération et de Développement Économiques"),
        ("CCSPPP-BAD","Cellule de Coordination et de Suivi du Portefeuille des Projets de la BAD"),
        ("CDC",       "Cahier des Charges"),
        ("CEDEAO",    "Communauté Économique des États de l'Afrique de l'Ouest"),
        ("CRUD",      "Create, Read, Update, Delete (opérations de base sur les données)"),
        ("DGCOD",     "Direction Générale de la Coopération au Développement"),
        ("DGP",       "Direction Générale du Plan"),
        ("DGATDRL",   "Direction Générale de l'Aménagement du Territoire et du Développement Régional et Local"),
        ("DirCab",    "Directeur de Cabinet"),
        ("DSID",      "Direction des Systèmes Informatiques et de la Digitalisation"),
        ("ENSEA",     "École Nationale Supérieure de Statistique et d'Économie Appliquée"),
        ("FAD",       "Fonds Africain de Développement (guichet concessionnel de la BAD)"),
        ("FADES",     "Fonds Arabe pour le Développement Économique et Social"),
        ("FIDA / IFAD","Fonds International de Développement Agricole"),
        ("FKDEA",     "Fonds Koweïtien pour le Développement Économique Arabe"),
        ("FSD",       "Fonds Saoudien pour le Développement"),
        ("ISDB / BIsD","Banque Islamique de Développement"),
        ("JICA",      "Japan International Cooperation Agency"),
        ("KPI",       "Key Performance Indicator — Indicateur Clé de Performance"),
        ("MPD",       "Ministère du Plan et du Développement"),
        ("OFID",      "Fonds de l'OPEP pour le Développement International"),
        ("PND",       "Plan National de Développement"),
        ("PTF",       "Partenaires Techniques et Financiers"),
        ("RBAC",      "Role-Based Access Control — Contrôle d'accès par rôles"),
        ("SYNAPSE",   "Système d'Information de la DGP pour le suivi des investissements publics"),
        ("UEMOA",     "Union Économique et Monétaire Ouest-Africaine"),
        ("XOF",       "Franc CFA de l'Afrique de l'Ouest — monnaie nationale"),
    ]
    big_table(doc, ["Sigle / Abréviation", "Signification"], abbrevs,
              col_widths=[4.5, 12], font_size=9.5)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — CONTEXTE
# ══════════════════════════════════════════════════════════════════════════════

def section1(doc):
    heading1(doc, "Contexte général et enjeux institutionnels", num="1.")
    separator(doc)

    heading2(doc, "Le Ministère du Plan et du Développement", num="1.1")
    body(doc,
        "Le Ministère du Plan et du Développement (MPD) est l'institution gouvernementale chargée de la "
        "conception, de la coordination et du suivi de la politique nationale de développement en Côte d'Ivoire. "
        "Il est responsable de l'élaboration du Plan National de Développement (PND), document stratégique "
        "pluriannuel qui fixe les priorités et les objectifs de développement du pays à moyen terme. "
        "À ce titre, il constitue l'interface institutionnelle centrale entre le gouvernement et l'ensemble "
        "des partenaires techniques et financiers (PTF) intervenant sur le territoire national.")
    body(doc,
        "Le Cabinet du Ministre constitue l'organe de pilotage stratégique et de coordination de haut niveau. "
        "Il assure la supervision du portefeuille global des projets de développement, coordonne les relations "
        "avec les bailleurs de fonds et veille à l'alignement des interventions extérieures sur les priorités "
        "nationales. Le Directeur de Cabinet joue un rôle d'interface opérationnel entre le Ministre et les "
        "directions techniques, assurant le suivi quotidien des dossiers et la remontée des alertes nécessitant "
        "un arbitrage ministériel.")
    body(doc,
        "Plusieurs directions générales et structures spécialisées opèrent sous la tutelle ou en lien direct "
        "avec le Ministère : la Direction Générale du Plan (DGP), la Direction Générale de la Coopération "
        "au Développement (DGCOD), la Direction Générale de l'Aménagement du Territoire et du Développement "
        "Régional et Local (DGATDRL), le Bureau National de la Prospective et de la Veille Stratégique (BNPVS), "
        "ainsi que les cellules de coordination des portefeuilles de chaque bailleur de fonds (cellule BAD, "
        "cellule Bailleurs Arabes, etc.). Ces structures constituent les Points Focaux techniques qui "
        "alimentent la plateforme en données opérationnelles.")

    heading2(doc, "Le rôle de gouverneur du Ministre et son portefeuille de bailleurs", num="1.2")
    body(doc,
        "Dans le cadre de ses fonctions, le Ministre du Plan et du Développement occupe la fonction de gouverneur "
        "de la Côte d'Ivoire auprès d'un ensemble d'institutions financières internationales et de fonds de "
        "développement multilatéraux. Cette fonction de gouverneur confère au Ministre une responsabilité "
        "directe dans la supervision des engagements contractés et des décaissements réalisés par ces "
        "institutions au bénéfice de la Côte d'Ivoire.")
    body(doc,
        "La présente plateforme a pour objet prioritaire de couvrir l'ensemble des projets co-financés par "
        "les bailleurs auprès desquels le Ministre détient cette responsabilité de gouverneur. Ce périmètre "
        "constitue la cible principale de la Phase 1, qui pourra être élargie dans les phases ultérieures "
        "à d'autres partenaires. La liste complète et définitive de ces bailleurs, devant être validée par "
        "le Cabinet, est présentée à l'Annexe A du présent document.")

    heading2(doc, "Les principaux bailleurs concernés par la Phase 1", num="1.3")
    body(doc,
        "Les projets à intégrer dans la plateforme lors de la Phase 1 sont ceux financés par les institutions "
        "et fonds de développement listés ci-après. Cette liste est indicative et sera arrêtée définitivement "
        "par le Cabinet du Ministère lors de la réunion de lancement du projet. Elle couvre les principaux "
        "partenaires multilatéraux, bilatéraux et les fonds de développement arabo-islamiques dont le "
        "Ministre est gouverneur ou alterno-gouverneur.")

    # Catégories de bailleurs
    heading3(doc, "a) Institutions de Bretton Woods et affiliées", num="")
    bailleurs_bw = [
        ("Banque mondiale — IDA / BIRD", "Association Internationale de Développement (guichet concessionnel) et Banque Internationale pour la Reconstruction et le Développement. La Côte d'Ivoire bénéficie d'un important portefeuille IDA. Le Ministre du Plan est typiquement alterno-gouverneur."),
        ("Société Financière Internationale (SFI / IFC)", "Guichet secteur privé du Groupe Banque mondiale."),
        ("Agence Multilatérale de Garantie des Investissements (MIGA)", "Garanties d'investissement — Groupe Banque mondiale."),
    ]
    big_table(doc, ["Institution", "Précisions"], bailleurs_bw, col_widths=[5, 11.5], font_size=9.5)

    heading3(doc, "b) Banques multilatérales de développement", num="")
    bailleurs_mdb = [
        ("Banque Africaine de Développement (BAD / ADB)", "Principal bailleur multilatéral africain, dont le siège est à Abidjan. Portfolio important sur tous les secteurs. Le Ministre du Plan est gouverneur ou alterno-gouverneur."),
        ("Fonds Africain de Développement (FAD)", "Guichet concessionnel de la BAD, financement des pays à faible revenu."),
        ("Banque Ouest Africaine de Développement (BOAD)", "Institution régionale de l'UEMOA, financements en FCFA."),
        ("Banque d'Investissement et de Développement de la CEDEAO (BIDC)", "Institution régionale CEDEAO."),
        ("Banque Européenne d'Investissement (BEI / EIB)", "Institution financière de l'Union Européenne."),
    ]
    big_table(doc, ["Institution", "Précisions"], bailleurs_mdb, col_widths=[5.5, 11], font_size=9.5)

    heading3(doc, "c) Fonds et institutions arabes et islamiques", num="")
    bailleurs_arabes = [
        ("Banque Islamique de Développement (BIsD / IsDB)", "Institution multilatérale à vocation islamique — important portefeuille en Côte d'Ivoire."),
        ("Fonds Koweïtien pour le Développement Économique Arabe (FKDEA)", "Fonds bilatéral koweïtien — financement de projets d'infrastructure."),
        ("Fonds Saoudien pour le Développement (FSD)", "Fonds bilatéral saoudien."),
        ("Fonds Arabe pour le Développement Économique et Social (FADES)", "Fonds arabe multilatéral."),
        ("Fonds de l'OPEP pour le Développement International (OFID)", "Fonds multilatéral des pays membres de l'OPEP."),
        ("Fonds Abu Dhabi pour le Développement (FADD)", "Fonds bilatéral des Émirats Arabes Unis."),
        ("Fonds d'Abu Dhabi pour le Développement — ADFD", "Autre guichet des EAU."),
    ]
    big_table(doc, ["Institution", "Précisions"], bailleurs_arabes, col_widths=[5.5, 11], font_size=9.5)

    heading3(doc, "d) Coopérations bilatérales et agences d'exécution", num="")
    bailleurs_bilat = [
        ("Agence Française de Développement (AFD)", "Principal bailleur bilatéral — large portefeuille en Côte d'Ivoire."),
        ("Proparco", "Filiale de l'AFD dédiée au secteur privé."),
        ("Japan International Cooperation Agency (JICA)", "Coopération japonaise — dons et prêts concessionnels."),
        ("Kreditanstalt für Wiederaufbau (KfW)", "Banque allemande de développement."),
        ("GIZ (Deutsche Gesellschaft für Internationale Zusammenarbeit)", "Coopération technique allemande."),
        ("US Millennium Challenge Corporation (MCC)", "Programme compact américain."),
        ("USAID", "Agence américaine pour le développement international."),
        ("Union Européenne (UE / FED)", "Fonds Européen de Développement et instruments de l'UE."),
        ("Exim Bank of China", "Financements chinois — projets d'infrastructure."),
        ("Fonds International de Développement Agricole (FIDA / IFAD)", "Institution du système des Nations Unies — focus agriculture."),
        ("Agence Coréenne de Coopération Internationale (KOICA)", "Coopération technique et financière de la Corée du Sud."),
    ]
    big_table(doc, ["Institution / Agence", "Précisions"], bailleurs_bilat, col_widths=[5.5, 11], font_size=9.5)

    heading3(doc, "e) Système des Nations Unies (agences co-finançant des projets)", num="")
    bailleurs_onu = [
        ("Programme des Nations Unies pour le Développement (PNUD)", "Appuis au développement humain, gouvernance, résilience."),
        ("Fonds des Nations Unies pour l'Enfance (UNICEF)", "Projets protection de l'enfance, éducation, santé."),
        ("Organisation Mondiale de la Santé (OMS / WHO)", "Projets santé."),
        ("Programme Alimentaire Mondial (PAM / WFP)", "Sécurité alimentaire."),
        ("FAO", "Agriculture et sécurité alimentaire."),
    ]
    big_table(doc, ["Agence ONU", "Domaine d'intervention"], bailleurs_onu, col_widths=[5.5, 11], font_size=9.5)

    info_box(doc,
        "📋  Note de périmètre :",
        "La liste ci-dessus est indicative. Elle sera finalisée et validée par le Cabinet du Ministère "
        "lors de l'atelier de lancement du projet, en concertation avec les Directions Générales concernées. "
        "Le prestataire devra prévoir une architecture permettant d'ajouter de nouveaux bailleurs sans "
        "modification structurelle de la plateforme.", bg='E8F5E9', border_color='009A44', title_color=VERT_CI)

    heading2(doc, "Diagnostic du dispositif actuel de suivi", num="1.4")
    body(doc,
        "Avant la mise en place de la présente plateforme, le suivi des projets des partenaires techniques "
        "et financiers au sein du Cabinet du Ministère repose principalement sur des outils bureautiques "
        "— essentiellement des tableurs Excel et des documents Word — disséminés entre les différentes "
        "cellules de coordination et directions techniques. Cette organisation génère plusieurs difficultés "
        "structurelles qui nuisent à l'efficacité du pilotage :")
    bullet(doc, "L'absence d'une base de données centralisée et unifiée empêche d'avoir une vision consolidée et en temps réel du portefeuille global. Chaque cellule dispose de ses propres tableaux de bord, dont les formats et les méthodes de calcul diffèrent, rendant toute agrégation laborieuse et sujette à erreurs.", bold_prefix="Fragmentation de l'information — ")
    bullet(doc, "La production d'indicateurs de synthèse pour le Ministre ou le Directeur de Cabinet nécessite de collecter manuellement les données auprès de chaque cellule, de les consolider et de les mettre en forme — un processus qui peut prendre plusieurs jours et qui génère inévitablement des décalages temporels importants.", bold_prefix="Délais de production des rapports — ")
    bullet(doc, "Les projets présentant des retards ou de faibles taux de décaissement ne sont pas automatiquement signalés. Leur identification requiert un examen manuel des tableaux de bord, ce qui ne permet pas une réactivité suffisante pour un pilotage efficace.", bold_prefix="Absence de système d'alerte — ")
    bullet(doc, "La traçabilité des données est insuffisante : il est difficile de savoir qui a modifié quoi et quand, ce qui peut poser des problèmes en cas de litige ou d'audit.", bold_prefix="Manque de traçabilité — ")
    bullet(doc, "L'absence d'un système structuré de gestion des droits d'accès ne permet pas de garantir que seuls les agents autorisés puissent modifier les données d'un projet donné.", bold_prefix="Sécurité et contrôle d'accès déficients — ")
    bullet(doc, "Les informations financières (engagements, décaissements, taux de conversion) ne sont pas toujours cohérentes d'une source à l'autre, faute d'une référence unique.", bold_prefix="Hétérogénéité des données financières — ")

    heading2(doc, "Démarche consultative et parties prenantes", num="1.5")
    body(doc,
        "La conception du présent cahier des charges s'est appuyée sur une démarche participative conduite "
        "du 5 au 8 mai 2026 sous l'égide du Cabinet du Ministère. Des séances de travail ont été organisées "
        "avec l'ensemble des structures concernées, à savoir la DSID, la Cellule de Coordination des Bailleurs "
        "Arabes, la CCSPPP-BAD, la DGP, la DGATDRL, la DGCOD, le BNPVS, l'ENSEA et le PHAS/ANStat. "
        "Ces consultations ont permis de recueillir les besoins opérationnels, les attentes spécifiques par profil "
        "d'utilisateur, ainsi qu'un ensemble de recommandations techniques et fonctionnelles.")
    body(doc,
        "Les principales orientations issues de cette démarche consultative, qui ont guidé la rédaction du "
        "présent cahier des charges, sont les suivantes : la nécessité de concevoir des interfaces différenciées "
        "selon le profil de l'utilisateur (Ministre, Directeur de Cabinet, Point Focal), la centralisation de "
        "la gestion des comptes, la priorité donnée à la simplicité d'utilisation et à la fiabilité des données, "
        "l'intégration d'un système d'alertes intelligent, et la nécessité de respecter la charte graphique du "
        "Ministère. Une synthèse complète des observations est présentée en Annexe B.")


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — OBJECTIFS ET VISION
# ══════════════════════════════════════════════════════════════════════════════

def section2(doc):
    heading1(doc, "Objectifs et vision stratégique de la plateforme", num="2.")
    separator(doc)

    heading2(doc, "Vision et finalité de la plateforme", num="2.1")
    body(doc,
        "La plateforme numérique de suivi des projets et des financements des PTF est conçue, en premier lieu, "
        "comme un outil de décision et d'alerte à destination du Ministre du Plan et du Directeur de Cabinet. "
        "Elle doit leur permettre, à tout moment, d'avoir une vue fiable et actualisée du portefeuille de projets, "
        "de détecter les situations nécessitant leur arbitrage, et de disposer d'éléments factuels pour leurs "
        "prises de position dans les discussions avec les partenaires.")
    body(doc,
        "En second lieu, la plateforme est un outil de travail quotidien pour les cellules techniques et "
        "les Points Focaux, qui l'utilisent pour saisir et mettre à jour les données, signaler les difficultés "
        "et produire les rapports périodiques. Elle constitue ainsi le canal formel par lequel les informations "
        "opérationnelles remontent vers le Cabinet, garantissant la cohérence et la traçabilité des données.")
    body(doc,
        "À terme, la plateforme a vocation à devenir un système d'information de référence pour la gestion "
        "de l'aide publique au développement en Côte d'Ivoire, susceptible d'être connecté aux autres "
        "systèmes d'information existants (SYNAPSE, plateformes des cellules bailleurs) et, le cas échéant, "
        "d'être rendu accessible au public dans une version simplifiée.")

    heading2(doc, "Objectifs opérationnels", num="2.2")
    body(doc, "La plateforme devra atteindre les objectifs opérationnels suivants :")
    numbered_item(doc, "1",
        "Centraliser dans une base de données unique et sécurisée l'ensemble des informations relatives "
        "aux projets, bailleurs, financements et décaissements du portefeuille PTF du Ministère.")
    numbered_item(doc, "2",
        "Offrir des tableaux de bord différenciés selon le profil de l'utilisateur, allant d'une vue "
        "très synthétique pour le Ministre à une vue opérationnelle détaillée pour les Points Focaux.")
    numbered_item(doc, "3",
        "Détecter automatiquement les situations d'alerte (projets en retard, faibles taux de "
        "décaissement, discordances entre avancement physique et financier) et les signaler aux "
        "niveaux hiérarchiques appropriés.")
    numbered_item(doc, "4",
        "Permettre un suivi précis et fiable des flux financiers : montants engagés, décaissements "
        "effectués, reste à décaisser, taux de décaissement par projet et par bailleur.")
    numbered_item(doc, "5",
        "Offrir une visualisation cartographique des zones d'intervention des projets sur le "
        "territoire national, permettant d'identifier les régions bénéficiaires et les zones "
        "sous-couvertes.")
    numbered_item(doc, "6",
        "Faciliter la production de rapports et d'exports de données, notamment la liste des projets "
        "et les synthèses KPI, exportables par le Ministre et le Directeur de Cabinet.")
    numbered_item(doc, "7",
        "Permettre l'alimentation efficace de la base de données par import de fichiers Excel, "
        "pour les phases d'initialisation et de mise à jour périodique.")
    numbered_item(doc, "8",
        "Garantir la sécurité des données, la traçabilité des actions et la conformité avec "
        "la réglementation sur la protection des données personnelles.")

    heading2(doc, "Principes directeurs de conception", num="2.3")
    body(doc, "La conception de la plateforme devra être guidée par les principes suivants :")
    bullet(doc,
        "La plateforme est avant tout un outil d'alerte et de décision pour le Ministre et le Directeur de Cabinet. "
        "La priorité doit être donnée à la clarté, à la synthèse et à l'actionabilité de l'information présentée à "
        "ces profils. La complexité opérationnelle doit être absorbée par les interfaces des Points Focaux, "
        "invisibles pour la hiérarchie.", bold_prefix="Décision d'abord — ")
    bullet(doc,
        "Chaque profil d'utilisateur doit avoir accès à une interface adaptée à ses besoins et à son niveau "
        "de responsabilité, sans être submergé par des informations non pertinentes.", bold_prefix="Interface adaptée au profil — ")
    bullet(doc,
        "Les mécanismes de saisie doivent être conçus pour minimiser les risques d'erreur : menus déroulants, "
        "validations en temps réel, contrôle de cohérence, droits de modification limités à la structure "
        "responsable de chaque projet.", bold_prefix="Fiabilité des données — ")
    bullet(doc,
        "La charte graphique du Ministère du Plan et du Développement devra être strictement respectée "
        "dans la conception de l'interface (couleurs institutionnelles, typographie, logo).", bold_prefix="Charte graphique du Ministère — ")
    bullet(doc,
        "L'accès à la plateforme est strictement réservé aux agents autorisés. La gestion des comptes "
        "est centralisée au niveau de l'administrateur. Aucun mécanisme d'auto-inscription publique n'est prévu.", bold_prefix="Sécurité et contrôle des accès — ")
    bullet(doc,
        "La plateforme devra être conçue dès le départ pour pouvoir évoluer : ajout de nouveaux modules, "
        "intégration avec d'autres systèmes, extension du périmètre des bailleurs.", bold_prefix="Évolutivité — ")


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — PÉRIMÈTRE
# ══════════════════════════════════════════════════════════════════════════════

def section3(doc):
    heading1(doc, "Périmètre fonctionnel de la Phase 1", num="3.")
    separator(doc)

    heading2(doc, "Définition du périmètre", num="3.1")
    body(doc,
        "La Phase 1 de la plateforme couvre l'ensemble du cycle de vie des projets de développement "
        "financés par les PTF listés à l'Annexe A, depuis leur identification jusqu'à leur clôture. "
        "Elle intègre les modules fonctionnels listés dans le tableau ci-dessous.")
    modules = [
        ("M1", "Tableau de bord analytique", "Indicateurs KPI, graphiques de synthèse, alertes actives. Interfaces différenciées par profil."),
        ("M2", "Gestion des projets et programmes", "Référentiel complet des projets avec cycle de vie, indicateurs financiers et physiques, cartographie."),
        ("M3", "Gestion des bailleurs de fonds", "Répertoire des PTF avec fiche analytique, logos, portefeuille de projets."),
        ("M4", "Financements et décaissements", "Suivi détaillé des engagements et décaissements par projet, par bailleur et par type de financement."),
        ("M5", "Système d'alertes et de signalement", "Détection automatique des retards, faibles décaissements, discordances physique/financier."),
        ("M6", "Module cartographique", "Carte interactive des zones d'intervention, niveaux géographiques, fiche cartographique par projet."),
        ("M7", "Reporting et export", "Export Excel des projets et synthèses KPI, rapport bailleur, rapport retards."),
        ("M8", "Import de données en masse", "Import Excel structuré avec prévisualisation et rapport d'erreurs."),
        ("M9", "Assistant d'analyse IA", "Interface conversationnelle d'interrogation de la base de données par intelligence artificielle."),
        ("M10", "Administration des comptes et des droits", "Gestion centralisée des utilisateurs, rôles, permissions, journal d'audit."),
    ]
    big_table(doc, ["Réf.", "Module", "Description"], modules, col_widths=[1.2, 4.5, 10.8], font_size=9.5)

    heading2(doc, "Ce qui est hors périmètre (Phase 2 et au-delà)", num="3.2")
    body(doc,
        "Afin de maîtriser la complexité de la Phase 1 et de garantir une livraison dans les délais, "
        "les fonctionnalités suivantes sont expressément exclues du périmètre et reportées à une Phase 2 :")
    bullet(doc, "Alignement avec les piliers et sous-objectifs du Plan National de Développement (PND) : la mise en œuvre de cette fonctionnalité nécessite un travail préalable de spécification détaillée et de structuration du PND en cours. Elle sera intégrée dans la Phase 2.")
    bullet(doc, "Interconnexion avec le système SYNAPSE de la DGP : la faisabilité technique de cette interconnexion doit faire l'objet d'une étude préalable impliquant la DSID. Elle est exclue de la Phase 1.")
    bullet(doc, "Interface publique / portail citoyen : la plateforme Phase 1 est strictement à usage interne du Ministère.")
    bullet(doc, "Génération automatique de rapports PDF : cette fonctionnalité pourra être ajoutée en Phase 2.")
    bullet(doc, "Notifications par email ou SMS : reportées en Phase 2.")
    bullet(doc, "Gestion des marchés et contrats associés aux projets : hors périmètre Phase 1.")
    bullet(doc, "Module de planification des activités et sous-activités : hors périmètre Phase 1.")


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — PROFILS UTILISATEURS
# ══════════════════════════════════════════════════════════════════════════════

def section4(doc):
    heading1(doc, "Profils utilisateurs et interfaces dédiées", num="4.")
    separator(doc)

    heading2(doc, "Vue d'ensemble du système de gestion des accès", num="4.1")
    body(doc,
        "La plateforme est un outil à usage strictement interne du Ministère du Plan et du Développement. "
        "L'accès est contrôlé par un système de gestion des droits basé sur les rôles (RBAC — Role-Based "
        "Access Control). Aucun mécanisme d'auto-inscription n'est prévu : seul l'administrateur de la "
        "plateforme est habilité à créer des comptes et à en transmettre les identifiants aux utilisateurs "
        "autorisés. Cette exigence, fortement soulignée lors des consultations, garantit que l'accès à "
        "la plateforme reste strictement maîtrisé.")
    body(doc,
        "La plateforme définit quatre niveaux d'accès correspondant aux profils fonctionnels des utilisateurs. "
        "Chaque profil bénéficie d'une interface et d'un niveau d'information adaptés à ses responsabilités. "
        "Le tableau suivant présente la matrice des profils :")
    profils = [
        ("Administrateur\n(Super Admin)", "Directeur de la DSID ou agent désigné", "Accès total à toutes les données et fonctionnalités. Création et gestion des comptes. Configuration de la plateforme. Accès au journal d'audit complet."),
        ("Directeur /\nHaute Fonction", "Ministre, DirCab, Conseillers, DG", "Accès en lecture et export sur toutes les données. Interface synthétique pour le Ministre. Interface exhaustive pour le DirCab et les conseillers."),
        ("Point Focal /\nCellule technique", "Points Focaux des cellules bailleurs, agents DGCOD/DGP", "Accès en saisie et modification restreint aux bailleurs/projets qui lui sont assignés. Accès en lecture aux autres données."),
        ("Lecteur", "Conseillers techniques, agents observateurs", "Accès en lecture seule sur les données du périmètre autorisé. Aucune modification possible."),
    ]
    big_table(doc, ["Profil", "Destinataires types", "Droits et périmètre"],
              profils, col_widths=[3.5, 4.5, 8.5], font_size=9.5)

    body(doc,
        "Les fonctions institutionnelles à responsabilité unique (Ministre, Directeur de Cabinet, etc.) "
        "ne peuvent être attribuées qu'à un seul compte actif à la fois. L'administrateur s'assure "
        "qu'aucune duplication ne soit possible. Il est rappelé que la plateforme devra intégrer "
        "les exigences de la réglementation ivoirienne en matière de protection des données personnelles "
        "(Loi n°2013-450 relative à la protection des données à caractère personnel et cadre ARTCI).")

    heading2(doc, "Interface Ministre — Vue synthétique de décision", num="4.2")
    body(doc,
        "Compte tenu de ses responsabilités et de la nature de son rôle, le Ministre dispose d'une interface "
        "spécifique, épurée et fortement synthétique, qui lui présente exclusivement les informations "
        "stratégiques dont il a besoin pour exercer ses responsabilités de gouvernance. Cette interface "
        "ne le sollicite que pour les sujets nécessitant effectivement son arbitrage.")
    profile_card(doc,
        "Interface Ministre",
        "Vue de décision stratégique — accès en lecture, alerte et export",
        [
            "Tableau de bord ultra-synthétique : 4 à 6 indicateurs clés en grand format (nombre de projets actifs, montant total engagé, taux de décaissement global, nombre de projets en retard, nombre de bailleurs actifs).",
            "Alertes stratégiques mises en évidence : projets en retard critique nécessitant un arbitrage, taux de décaissement anormalement faibles, alertes signalées par les cellules techniques.",
            "Graphiques simples et lisibles : répartition sectorielle, evolution des décaissements, top 5 bailleurs.",
            "Carte géographique des zones d'intervention : visualisation rapide de la couverture territoriale.",
            "Accès à la liste des projets avec filtre rapide et bouton d'export en un clic.",
            "Export de la vue synthétique des KPI (PDF ou Excel) pour réunions et conseils ministériels.",
            "Barre de navigation réduite à l'essentiel — pas d'accès aux interfaces de saisie.",
        ],
        bg_hex='FFF8E1', accent='F77F00'
    )

    heading2(doc, "Interface Directeur de Cabinet — Vue exhaustive de pilotage", num="4.3")
    body(doc,
        "Le Directeur de Cabinet est l'interface opérationnelle entre le Ministre et les structures techniques. "
        "Son interface est plus exhaustive que celle du Ministre, lui permettant d'accéder à l'ensemble "
        "des données du portefeuille, de suivre les remontées des Points Focaux et de préparer les dossiers "
        "à soumettre à l'arbitrage ministériel. Elle constitue le principal outil de pilotage quotidien.")
    profile_card(doc,
        "Interface Directeur de Cabinet",
        "Vue de pilotage opérationnel — accès en lecture, analyse et export complets",
        [
            "Tableau de bord complet avec tous les KPIs, graphiques détaillés, évolutions dans le temps.",
            "Accès au centre de notifications : projets en retard, faibles décaissements, activités récentes des Points Focaux, alertes signalées.",
            "Accès à la liste complète des projets avec filtres multicritères avancés (statut, secteur, bailleur, zone, retard).",
            "Fiche détaillée de chaque projet : informations complètes, financements, décaissements, motifs de retard.",
            "Fiche analytique par bailleur : portefeuille, indicateurs financiers, carte.",
            "Accès aux analyses interactives : moteur de graphiques personnalisables sur l'ensemble du portefeuille.",
            "Export de la liste des projets et des synthèses KPI.",
            "Accès à l'assistant IA pour interroger la base de données en langage naturel.",
            "Barre de navigation complète — sans accès à l'interface d'administration des comptes.",
        ],
        bg_hex='E8F4FB', accent='1A5E8A'
    )

    heading2(doc, "Interface Point Focal / Cellule technique — Vue opérationnelle", num="4.4")
    body(doc,
        "Les Points Focaux sont les agents chargés de la saisie et de la mise à jour des données "
        "de la plateforme. Chaque Point Focal est assigné à un ou plusieurs bailleurs spécifiques "
        "et ne peut modifier que les données relevant de son périmètre. Ce mécanisme, "
        "fortement demandé lors des consultations, garantit l'intégrité des données : "
        "une seule structure désignée est responsable des informations d'un projet donné, "
        "les autres n'ayant qu'un accès en lecture.")
    profile_card(doc,
        "Interface Point Focal / Cellule technique",
        "Vue opérationnelle — saisie et mise à jour des données du périmètre assigné",
        [
            "Tableau de bord personnel : état du portefeuille des bailleurs assignés, projets à mettre à jour, alertes actives.",
            "Accès en création et modification aux projets, financements et décaissements des bailleurs assignés uniquement.",
            "Accès en lecture sur les projets des autres bailleurs (sans modification possible).",
            "Formulaire de saisie des projets : données administratives, financières, avancement physique, motif de retard, responsable local, zone d'intervention.",
            "Saisie des décaissements : montant, date, référence, description.",
            "Import de données par fichier Excel (habilités désignés uniquement).",
            "Accès à l'assistant IA avec données filtrées sur son périmètre.",
            "Indicateur de ponctualité de saisie : suivi des délais de mise à jour des données.",
        ],
        bg_hex='E8F5E9', accent='009A44'
    )

    heading2(doc, "Interface Conseiller / Lecteur — Vue de consultation", num="4.5")
    body(doc,
        "Les conseillers techniques et autres agents bénéficiant d'un accès en lecture seule "
        "disposent d'une interface leur permettant de consulter les informations du portefeuille "
        "selon leur périmètre autorisé, sans possibilité de modification. Ce profil répond "
        "aux besoins de consultation ponctuels, notamment pour la préparation de notes techniques "
        "ou d'analyses sectorielles.")
    profile_card(doc,
        "Interface Lecteur / Conseiller",
        "Vue de consultation — accès en lecture seule",
        [
            "Accès en consultation à toutes les données du périmètre autorisé.",
            "Tableau de bord en lecture seule.",
            "Consultation des fiches projets et bailleurs.",
            "Accès à l'assistant IA pour des interrogations ponctuelles.",
            "Aucune possibilité de saisie, modification ou suppression.",
        ],
        bg_hex='F5F5F5', accent='7A8A9A'
    )


# ══════════════════════════════════════════════════════════════════════════════
# Écriture dans un fichier intermédiaire pour appel depuis gen_cdc_v2_part2.py
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    # Import de la seconde partie
    from gen_cdc_v2_part2 import section5, section6, section7, section8, annexes, build_all
    build_all()
